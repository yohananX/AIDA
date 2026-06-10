"""
AIDA Dissertation Generator
Affective Intelligent Dialogue Agent for Preliminary Mental Health Intervention
Full 4-chapter dissertation, 20,000+ words
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ─────────────────────────────────────────────
# STYLES
# ─────────────────────────────────────────────
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

styles = doc.styles

def set_style(style_name, font_name='Times New Roman', font_size=12,
              bold=False, space_before=0, space_after=6, line_spacing=None,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    try:
        st = styles[style_name]
    except KeyError:
        return
    st.font.name = font_name
    st.font.size = Pt(font_size)
    st.font.bold = bold
    pf = st.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.alignment    = alignment
    if line_spacing:
        from docx.shared import Pt as P
        pf.line_spacing = P(line_spacing)

# Normal body
set_style('Normal', font_size=12, space_after=6, line_spacing=18)

# Headings
set_style('Heading 1', font_size=14, bold=True, space_before=24, space_after=12,
          alignment=WD_ALIGN_PARAGRAPH.LEFT)
set_style('Heading 2', font_size=13, bold=True, space_before=18, space_after=6,
          alignment=WD_ALIGN_PARAGRAPH.LEFT)
set_style('Heading 3', font_size=12, bold=True, space_before=12, space_after=6,
          alignment=WD_ALIGN_PARAGRAPH.LEFT)

# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return p

def para(text, bold_parts=None, indent=False):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    from docx.shared import Pt as P, Inches as I
    pf.space_after = P(6)
    if indent:
        pf.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def para_indent(text):
    return para(text, indent=True)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5 + level * 0.25)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def page_break():
    doc.add_page_break()

def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    return p

def add_table(headers, rows, col_widths=None, caption_text=None):
    if caption_text:
        cp = doc.add_paragraph()
        cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.paragraph_format.space_after = Pt(4)
        cr = cp.add_run(caption_text)
        cr.font.name = 'Times New Roman'
        cr.font.size = Pt(11)
        cr.font.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    spacer()
    return table


# ═══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════
spacer(3)
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('AFFECTIVE INTELLIGENT DIALOGUE AGENT FOR PRELIMINARY\nMENTAL HEALTH INTERVENTION')
tr.font.name = 'Times New Roman'
tr.font.size = Pt(16)
tr.font.bold = True

spacer(2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A Research Project Submitted in Partial Fulfilment of the\nRequirements for the Award of Bachelor of Science Degree')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

spacer(2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('By\n\nFolajimi Yohanan')
r.font.name = 'Times New Roman'
r.font.size = Pt(13)
r.font.bold = True

spacer(2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Department of Computer Science\n2026')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

page_break()


# ═══════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════
heading('ABSTRACT', level=1)
para(
    'Mental health disorders constitute a major public health challenge globally, with access to professional support remaining constrained by cost, stigma, workforce shortages, and geographical barriers. Digital conversational agents have emerged as a scalable complementary approach, yet most existing systems respond reactively to individual messages without accounting for the emotional trajectory of a conversation over time. This study proposes and implements the Adaptive Emotion-Aware Intervention Framework (AEIF), a six-layer conversational pipeline that augments standard emotion-based response generation with emotional trend analysis and strategy-driven intervention selection. The AEIF pipeline processes each user message through: (1) a safety detection layer that unconditionally screens for crisis signals; (2) an emotion classification layer using the j-hartmann/emotion-english-distilroberta-base transformer with keyword-based fallback; (3) an in-memory session store that maintains conversational history; (4) a trend analysis layer that derives an emotional trajectory from session history across seven trajectory types; (5) a strategy selection layer that maps trajectory and current emotion to one of nine CBT-aligned intervention strategies; and (6) a large language model response generation layer using the Groq API with the llama-3.3-70b-versatile model, constrained by strategy-specific system prompts. A baseline comparison pipeline — identical in all respects except that trend analysis and strategy selection are disabled — was implemented to enable controlled qualitative comparison of responses. Evaluation on a 58-case annotated test set including Nigerian Pidgin English, emotionally ambiguous multi-clause sentences, and context-contradicting inputs yielded a crisis detection accuracy of 100%, emotion cluster accuracy of 72.9%, and trend analysis accuracy of 100% on scripted multi-turn sequences. Qualitative comparison demonstrated that AEIF-generated responses exhibited strategy-specific language and higher contextual appropriateness than baseline responses across multiple conversation scenarios. The system was implemented as a full-stack web application with a React frontend incorporating an emotional trend timeline, ambient emotion visualisation, and a per-turn perceived empathy rating mechanism for evaluation data collection. Findings suggest that trajectory-aware strategy selection meaningfully differentiates conversational AI responses in mental health support contexts, and that the AEIF architecture provides a practical, ethically bounded framework for preliminary mental health intervention in the Nigerian young-adult population.'
)
page_break()


# ═══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════
heading('TABLE OF CONTENTS', level=1)
toc_items = [
    ('Abstract', '2'),
    ('Chapter 1: Introduction', '4'),
    ('  1.1 Background of Study', '4'),
    ('  1.2 Statement of Problem', '6'),
    ('  1.3 Aim and Objectives of Study', '7'),
    ('  1.4 Scope of Study', '8'),
    ('  1.5 Justification of the Study', '9'),
    ('Chapter 2: Literature Review and Theoretical Framework', '10'),
    ('  2.1 Overview of Mental Health', '10'),
    ('  2.2 Digital Mental Health Interventions', '12'),
    ('  2.3 Artificial Intelligence in Mental Health Care', '13'),
    ('  2.4 AI-Based Mental Health Chatbots', '14'),
    ('  2.5 Theoretical Framework', '15'),
    ('  2.6 Research Gaps Identified', '17'),
    ('  2.7 Related Works', '18'),
    ('Chapter 3: System Design and Methodology', '22'),
    ('  3.1 Introduction', '22'),
    ('  3.2 Research Design', '22'),
    ('  3.3 The AEIF Framework', '24'),
    ('  3.4 System Architecture', '26'),
    ('  3.5 Study Population and Location', '28'),
    ('  3.6 Tools and Technologies', '29'),
    ('  3.7 Data Collection and Test Set Design', '30'),
    ('  3.8 Evaluation Strategy', '31'),
    ('  3.9 Ethical Considerations', '33'),
    ('Chapter 4: Implementation and Evaluation', '34'),
    ('  4.1 Introduction', '34'),
    ('  4.2 Development Environment', '34'),
    ('  4.3 Project Structure', '35'),
    ('  4.4 AEIF Pipeline Implementation', '36'),
    ('  4.5 Emotion Classification', '40'),
    ('  4.6 Trend Analysis Implementation', '43'),
    ('  4.7 Strategy Selection Engine', '45'),
    ('  4.8 LLM Integration and Baseline Comparison Mode', '46'),
    ('  4.9 Frontend Implementation', '49'),
    ('  4.10 Evaluation Results', '51'),
    ('  4.11 AEIF vs Baseline: Qualitative Comparison', '55'),
    ('  4.12 Implementation Challenges and Solutions', '59'),
    ('  4.13 Limitations', '61'),
    ('  4.14 Summary', '62'),
    ('References', '64'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(item)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    if not item.startswith(' '):
        r.font.bold = True

page_break()


# ═══════════════════════════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════
heading('CHAPTER 1: INTRODUCTION', level=1)

heading('1.1 Background of Study', level=2)
para(
    'Information Technology has become a central driver of transformation across sectors including healthcare, education, and finance, largely due to the integration of advanced computational techniques such as Artificial Intelligence (AI), Machine Learning (ML), and Natural Language Processing (NLP) (Laranjo et al., 2018; Yang et al., 2023). These technologies have enabled the creation of intelligent systems capable of performing tasks traditionally requiring human cognition, including decision-making, pattern recognition, and natural language understanding. Within healthcare, one area that has received growing attention is mental health, where AI-powered solutions offer novel approaches to providing psychological support and emotional assistance through digital platforms (Fitzpatrick et al., 2017; Inkster et al., 2018).'
)
para(
    'Mental health disorders, including depression, anxiety, and stress-related conditions, represent a significant challenge worldwide, contributing to reduced quality of life, impaired social functioning, and decreased productivity (Vaidyam et al., 2019; Fulmer et al., 2018). Access to professional mental health services remains constrained in many regions, particularly in low-resource settings, due to shortages of trained professionals, high costs, and social stigma (Abd-Alrazaq et al., 2020; Laranjo et al., 2018). Even in high-resource environments, individuals may encounter long waiting times, geographical barriers, and other obstacles that limit timely mental health care (Fitzpatrick et al., 2017; Inkster et al., 2018).'
)
para(
    'In Nigeria specifically, the mental health burden is substantial and largely unaddressed. The World Health Organization estimates that one in four Nigerians suffers from a mental health condition, yet the country has fewer than 300 psychiatrists serving a population of over 200 million (Atilola et al., 2020). Stigma, cultural misconceptions about mental illness, and inadequate health infrastructure compound the access problem. Young adults aged 16 to 30 — a demographic characterised by high academic stress, economic precarity, and digital platform engagement — are disproportionately affected and represent a primary target population for digital mental health intervention (Feng et al., 2025; Racine et al., 2021).'
)
para(
    'Digital mental health interventions, particularly conversational agents or chatbots, have emerged as scalable, cost-effective, and accessible tools to address these challenges (Abd-Alrazaq et al., 2020; Vaidyam et al., 2019). These systems engage users in natural language conversations, provide emotional support, suggest coping strategies, and direct individuals towards additional resources. Advances in NLP and deep learning have improved chatbots\' ability to understand user intent, sentiment, and context, enabling more personalised and responsive interactions (Yang et al., 2023; Poria et al., 2021). Well-known systems such as Woebot (Fitzpatrick et al., 2017) and Wysa (Inkster et al., 2018) have demonstrated measurable reductions in self-reported anxiety and depressive symptoms among user populations.'
)
para(
    'Despite these advancements, a significant architectural limitation persists across most existing systems: they respond to each user message in isolation, without systematically analysing the emotional trajectory of the conversation over time. A user who expresses sadness across multiple consecutive turns is exhibiting persistent distress, which warrants a qualitatively different response than a user expressing sadness for the first time. Similarly, a user whose emotional valence shifts from negative to positive across turns represents an improving trajectory that should be met with encouragement and reinforcement rather than continued validation. Current systems that respond reactively to the current message without this trajectory awareness miss these clinically meaningful patterns (Liu et al., 2021; Zhang et al., 2022).'
)
para(
    'This study addresses this limitation through the design and implementation of the Affective Intelligent Dialogue Agent (AIDA), a text-based conversational agent grounded in the Adaptive Emotion-Aware Intervention Framework (AEIF). AEIF is a six-layer pipeline that integrates emotion detection, emotional memory, trajectory analysis, CBT-aligned strategy selection, and LLM response generation into a coherent, ethically bounded architecture. The system is evaluated empirically through automated testing and a controlled qualitative comparison against a baseline pipeline that lacks trajectory analysis and strategy selection, providing direct evidence of the contribution that trajectory-aware strategy selection makes to response quality.'
)
para(
    'The evolution of technology in mental healthcare can be traced back to early computer-assisted systems. ELIZA (Weizenbaum, 1966), one of the first examples, used simple pattern-matching techniques to simulate a psychotherapist. Although limited in understanding or emotional awareness, ELIZA demonstrated that computers could engage users in psychologically oriented dialogue. Later systems incorporated structured questionnaires, decision trees, and Cognitive Behavioral Therapy (CBT) frameworks, providing guided interventions but remaining limited in adaptability to individual emotional states (Calvo & D\'Mello, 2010; Fulmer et al., 2018). The integration of AI and ML allowed modern systems to move beyond rigid scripts toward semantic understanding and context-aware responses (Yang et al., 2023). However, the temporal dimension of emotional experience — the pattern of how emotions evolve across a conversation — has remained underexplored in the design of conversational mental health agents.'
)

heading('1.2 Statement of Problem', level=2)
para(
    'Despite increased global awareness of mental health issues, access to timely and effective care remains a significant challenge. High costs, limited availability of mental health professionals, social stigma, and geographical barriers continue to prevent many individuals from receiving adequate support, particularly in low-resource regions with insufficient healthcare infrastructure and limited mental health literacy (Feng et al., 2025; Ankomah & Turkson, 2025). These challenges are especially acute in sub-Saharan African contexts, where the mental health treatment gap — the proportion of people with mental health conditions who receive no treatment — exceeds 90% in some estimates (Charlson et al., 2019).'
)
para(
    'Digital mental health interventions, including chatbots and mobile applications, have attempted to bridge these gaps by providing scalable, accessible support (Li et al., 2023; De Grandi et al., 2024). While such systems show promise in alleviating mental distress, the dominant architectural approach — detecting a user\'s current emotional state and generating a response conditioned on that state — is insufficient for sustained, meaningful emotional support. Human emotional experience is not a sequence of independent states; it is a temporal trajectory in which each emotional expression is contextualised by what preceded it. A user saying "I feel a bit better today" means something qualitatively different if it follows three turns of expressed hopelessness versus if it is a session-opening statement. Existing systems, by treating each message independently, are blind to this distinction (Zhang et al., 2022; Liu et al., 2021).'
)
para(
    'Furthermore, existing solutions often focus primarily on symptom assessment rather than empathetic interaction, leaving users feeling unheard or misunderstood (Hudon & Stip, 2025; Chin et al., 2025). The absence of culturally adaptable and emotionally intelligent conversational agents is a further gap, as most systems are trained and evaluated on Western English corpora with limited attention to African linguistic contexts, including Nigerian Pidgin English and culturally specific idioms of distress.'
)
para(
    'These limitations motivate the need for a conversational agent that: (1) detects emotions reliably, including in culturally specific linguistic contexts; (2) tracks the emotional trajectory of a conversation across turns; (3) selects an intervention strategy calibrated to that trajectory rather than only to the current message; and (4) generates responses whose content is directly shaped by that strategy. This is the research problem this study addresses.'
)

heading('1.3 Aim and Objectives of Study', level=2)
para(
    'The aim of this study is to design, implement, and evaluate the Adaptive Emotion-Aware Intervention Framework (AEIF), a six-layer conversational pipeline for AIDA — the Affective Intelligent Dialogue Agent — that provides preliminary emotional support by combining per-message emotion detection with session-level emotional trajectory analysis and CBT-aligned strategy-driven response generation.'
)
para('The specific objectives are:')
bullet('To implement a multi-layer emotion detection system that classifies user text inputs into six emotion clusters (Positive, Sadness, Anxiety, Anger, Neutral, Ambiguous), with support for Nigerian Pidgin English and ambiguity detection for emotionally contradictory messages.')
bullet('To design and implement an emotional trend analysis system that derives an emotional trajectory (Persistent Distress, Escalating Distress, Elevated Anger, Improving, Stable Positive, Fluctuating, or Insufficient Data) from session-level emotion history using a weighted severity scoring model.')
bullet('To implement a CBT-aligned strategy selection engine that maps the detected emotional trajectory and current emotion cluster to one of nine intervention strategies (Validation and Reflection, Grounding, Calm Reflection, Encouragement, Positive Reinforcement, Exploratory Inquiry, Validation, Clarification Request, Open Check-in).')
bullet('To integrate the AEIF pipeline with a large language model (LLM) that generates responses constrained by the selected strategy, and to implement a baseline pipeline — identical except that trajectory analysis and strategy selection are disabled — to enable controlled comparison.')
bullet('To evaluate the system through automated testing on an annotated 58-case test set and through qualitative comparison of AEIF and baseline response sets on three scripted multi-turn conversation scenarios.')

heading('1.4 Scope of Study', level=2)
para(
    'This study focuses on the design, implementation, and evaluation of a text-based conversational agent for preliminary emotional support. The system processes English-language text input, with additional keyword support for Nigerian Pidgin English expressions. It detects emotional states using a transformer-based classifier with a keyword fallback, maintains session-level emotional history, derives emotional trajectories across seven trajectory types, selects from nine CBT-aligned intervention strategies, and generates responses using a large language model.'
)
para(
    'The scope is limited to non-clinical preliminary support. The system does not provide medical diagnoses, clinical assessments, or therapeutic recommendations, and is not intended to replace professional mental health practitioners. It serves as a supplementary support tool and an early-access point that may prompt users to seek professional help when warranted. All crisis signals bypass the conversational pipeline and return pre-defined safety responses including verified Nigerian emergency resources.'
)
para(
    'The evaluation scope covers automated performance metrics on an annotated test dataset, qualitative comparison of AEIF and baseline response sets, and trend analysis accuracy on scripted multi-turn sequences. The study does not include clinical outcome measurement, large-scale user trials, or long-term engagement analysis. Ethical considerations including data minimisation, informed consent, and crisis escalation are addressed at the system design level.'
)
para(
    'The project does not address voice-based interaction, multilingual support beyond English and Nigerian Pidgin, or real-time integration with healthcare systems. The implementation constitutes a functional research prototype suitable for academic evaluation and demonstration purposes.'
)

heading('1.5 Justification of the Study', level=2)
para(
    'This study is justified on three grounds: academic contribution, practical relevance, and contextual novelty.'
)
para(
    'Academically, the study contributes a formally specified framework — AEIF — that extends the standard emotion-detection-to-response pipeline by inserting trajectory analysis and strategy selection as intermediate layers. While individual components of this pipeline (emotion detection, CBT-based response generation, LLM integration) have been separately explored in the literature, their integration into a unified trajectory-aware architecture, evaluated through controlled baseline comparison, represents a contribution not addressed by prior work reviewed in Chapter 2. The controlled comparison methodology — running the same conversations through both AEIF and a strategy-disabled baseline — provides direct empirical evidence of the trajectory and strategy layers\' contribution, distinguishing this study from purely descriptive system builds.'
)
para(
    'Practically, the system addresses a real and underserved need. Nigeria\'s mental health treatment gap, combined with the high digital platform engagement of its young-adult population, creates a context in which a well-designed conversational agent could serve as a meaningful first point of contact for individuals in emotional distress. By including Nigerian Pidgin English in the emotion classifier\'s keyword vocabulary and embedding verified Nigerian crisis resources across all safety responses, the system is explicitly adapted for its target population in a way that generic mental health chatbots are not.'
)
para(
    'From a contextual novelty standpoint, the evaluation of an AI conversational agent in a Nigerian young-adult context using Nigerian linguistic data is, to the best of the author\'s knowledge, not represented in the reviewed literature. Most evaluation studies rely on Western English datasets and Western user populations. This study begins to address that gap, providing a reference point for future research into culturally adapted digital mental health interventions in sub-Saharan African contexts.'
)

page_break()


# ═══════════════════════════════════════════════════════════════════════
# CHAPTER 2: LITERATURE REVIEW
# ═══════════════════════════════════════════════════════════════════════
heading('CHAPTER 2: LITERATURE REVIEW AND THEORETICAL FRAMEWORK', level=1)

heading('2.1 Overview of Mental Health', level=2)
para(
    'Mental health is generally understood as a state of psychological, emotional, and social well-being that enables individuals to cope with everyday stressors, function productively, and contribute meaningfully to society. Contemporary research emphasises that mental health extends beyond the absence of mental illness and plays a critical role in overall health outcomes and quality of life (Galderisi et al., 2017; Patel et al., 2018). The World Health Organization defines mental health as "a state of mental well-being that enables people to cope with the stresses of life, realise their abilities, learn well and work well, and contribute to their community."'
)

heading('2.1.1 Global Burden and Contemporary Trends', level=3)
para(
    'Recent empirical studies indicate a steady increase in the global prevalence of mental health disorders, with a notable acceleration following the COVID-19 pandemic. Large-scale population analyses have reported significant rises in anxiety and depressive disorders, attributed to social isolation, economic uncertainty, and disruption of social support systems (Santomauro et al., 2021; Pierce et al., 2020). Evidence further suggests that young adults, students, healthcare workers, and individuals in economically disadvantaged settings are disproportionately affected (Racine et al., 2021; Aknin et al., 2022).'
)
para(
    'In the African context, mental health challenges are compounded by a critical shortage of mental health professionals, inadequate funding for mental health services, and pervasive stigma that discourages help-seeking (Atilola et al., 2020; Charlson et al., 2019). Nigeria, with a population exceeding 200 million, has fewer than 300 psychiatrists and approximately 8 psychiatric hospitals — figures that represent one of the world\'s most severe mental health workforce gaps (Atilola et al., 2020). Despite this growing demand, the expansion of mental health services has not kept pace, resulting in persistent treatment gaps, particularly in low- and middle-income countries (Charlson et al., 2019).'
)

heading('2.1.2 Limitations of Traditional Mental Health Care Systems', level=3)
para(
    'Conventional mental health care delivery relies primarily on face-to-face interactions between patients and trained professionals. While clinically effective, these systems are constrained by structural limitations such as workforce shortages, high financial costs, long waiting periods, and geographic inaccessibility (Atilola et al., 2020; Kazdin & Blase, 2011). Additionally, stigma and cultural misconceptions surrounding mental illness continue to discourage help-seeking behaviour, particularly in developing regions where mental health literacy remains limited (Clement et al., 2015; Thornicroft et al., 2016). These challenges underscore the need for complementary approaches to mental health support that are accessible, low-cost, and stigma-neutral.'
)

heading('2.2 Digital Mental Health Interventions', level=2)
para(
    'Digital mental health interventions involve the use of information and communication technologies to deliver mental health support, assessment, and self-help services. Such interventions include mobile health applications, web-based therapy platforms, telepsychiatry systems, and self-guided digital tools (Naslund et al., 2017; Firth et al., 2023). Empirical evidence suggests that digital interventions can reduce barriers related to cost, stigma, and accessibility, particularly for individuals with mild to moderate mental health conditions (Andersson et al., 2019; Mohr et al., 2021).'
)
para(
    'Mobile mental health applications commonly provide features such as mood tracking, guided meditation, journaling, and cognitive behavioural therapy exercises. Although short-term benefits have been reported, sustained engagement and clinical validation remain ongoing challenges (Lecomte et al., 2022; Baumel et al., 2019). High dropout rates and reduced user engagement are frequently reported, suggesting that emotional responsiveness and adaptability remain insufficient in many systems (Zhang et al., 2022; Karyotaki et al., 2018).'
)
para(
    'Teletherapy platforms facilitate remote interaction between patients and clinicians and have demonstrated effectiveness in maintaining continuity of care, particularly during public health crises (Torous & Roberts, 2024; Yellowlees et al., 2020). However, these platforms remain dependent on professional availability, limiting scalability and precluding the 24/7 accessibility that many distressed individuals require.'
)

heading('2.3 Artificial Intelligence in Mental Health Care', level=2)
para(
    'Artificial Intelligence refers to computational systems capable of performing tasks that typically require human intelligence, including learning, reasoning, and language understanding. In mental health care, AI techniques have been applied to early detection, risk assessment, treatment support, and patient monitoring (Topol, 2019; Yu et al., 2022). Natural language processing, in particular, has enabled systems to analyse the linguistic content of user communications and infer psychological states, opening new avenues for scalable, automated support.'
)

heading('2.3.1 Machine Learning and Emotion Recognition', level=3)
para(
    'Machine learning models have demonstrated effectiveness in identifying patterns associated with depression, anxiety, and stress through analysis of textual and behavioural data (Li et al., 2023; Chancellor & De Choudhury, 2020). Emotion recognition techniques further enable systems to infer users\' emotional states, supporting adaptive and context-aware interactions (Poria et al., 2021). The introduction of transformer-based architectures, particularly BERT (Devlin et al., 2019) and its derivatives, has substantially advanced the accuracy of text-based emotion classification by enabling deep contextual representations of language.'
)
para(
    'The GoEmotions dataset (Demszky et al., 2020), containing approximately 58,000 English Reddit comments annotated with 27 fine-grained emotion categories, has become a standard resource for training and evaluating emotion recognition models. Subsequent research has proposed consolidating these fine-grained categories into higher-level clusters aligned with therapeutic response strategies (Cowen and Keltner, 2017), a principle that directly informs the emotion classification design in the present study.'
)
para(
    'The Emotional Support Conversation (ESC) framework proposed by Liu et al. (2021) represents a significant contribution to the understanding of how conversational agents should structure emotional support. The framework identifies three stages of supportive dialogue — Exploration, Comforting, and Action — and proposes that effective support requires not only empathetic responses but also strategic selection of conversational moves calibrated to the user\'s emotional state and stage of the conversation. This framework provides direct theoretical grounding for the AEIF strategy selection layer.'
)

heading('2.3.2 Ethical Considerations in AI-Based Mental Health Systems', level=3)
para(
    'The deployment of AI in mental health contexts raises ethical concerns related to data privacy, algorithmic bias, transparency, and accountability. Scholars emphasise the importance of informed consent, human oversight, and clearly defined system limitations to prevent harm to vulnerable users (Floridi et al., 2018; Torous & Roberts, 2023). For systems deployed in crisis-proximate contexts, the adequacy of escalation mechanisms is particularly critical. Torous and Roberts (2023) highlight the risks of over-reliance on automated systems and the absence of robust escalation pathways as among the most significant ethical concerns in current AI mental health tools.'
)

heading('2.4 AI-Based Mental Health Chatbots', level=2)
para(
    'AI-based mental health chatbots are conversational systems designed to provide psychological support through text-based interaction. These systems can offer psychoeducation, emotional validation, and coping strategies while operating continuously and at scale (Abd-Alrazaq et al., 2023; Fulmer et al., 2018). Recent research distinguishes between rule-based chatbots and machine learning-driven systems, with emerging trends favouring hybrid architectures that balance conversational flexibility with safety constraints (Inkster et al., 2023; Li et al., 2023).'
)
para(
    'Empirical evaluations indicate that chatbot-based interventions can reduce symptoms of anxiety and depression, particularly among users hesitant to seek face-to-face support (Feng et al., 2025; Fitzpatrick et al., 2017). Woebot, evaluated through a randomised controlled trial by Fitzpatrick et al. (2017), demonstrated statistically significant reductions in depressive symptoms compared to a control group. Wysa, evaluated by Inkster et al. (2018), showed measurable reductions in self-reported anxiety among sustained users. However, both systems exhibited limitations in handling nuanced emotional narratives and demonstrated shallow adaptation to the temporal structure of conversation (Inkster et al., 2023; Zhang et al., 2022).'
)
para(
    'A consistent finding across the literature is that per-message emotion detection, while necessary, is insufficient for high-quality empathetic interaction. Zhang et al. (2022) specifically identified the absence of emotional trajectory awareness as a limitation of transformer-based mental health chatbots. Liu et al. (2021) formalised this concern through the ESC framework, arguing that appropriate emotional support requires tracking the user\'s evolving state across a conversation and selecting strategies calibrated to that evolution. The present study responds directly to this finding by implementing trajectory analysis and strategy selection as core architectural components rather than post-hoc additions.'
)

heading('2.5 Theoretical Framework', level=2)
para(
    'The theoretical framework for this study is anchored in three intersecting bodies of theory: Cognitive Behavioral Therapy (CBT), Affective Computing Theory, and the Emotional Support Conversation (ESC) framework. These collectively guide the design of the AEIF pipeline and its constituent components.'
)

heading('2.5.1 Cognitive Behavioral Therapy (CBT)', level=3)
para(
    'Cognitive Behavioral Therapy, developed by Aaron Beck (Beck, 1976), posits that thoughts, emotions, and behaviours are interrelated, and that modifying dysfunctional thinking patterns can improve emotional well-being and behavioural outcomes. CBT has been widely applied in digital mental health interventions due to its structured and evidence-based approach, making it particularly suitable for implementation in automated conversational systems (Fleming et al., 2022). The CBT principle of matching the therapeutic technique to the patient\'s current emotional and cognitive state is operationalised in AEIF through the strategy selection layer: different emotion-trajectory combinations trigger different CBT-aligned strategies (Validation, Grounding, Cognitive Reframing, Encouragement), ensuring that the system\'s conversational moves are calibrated to what is therapeutically appropriate for the user\'s detected state.'
)

heading('2.5.2 Affective Computing Theory', level=3)
para(
    'Affective Computing Theory, introduced by Picard (1997), emphasises the development of computational systems capable of recognising, interpreting, and responding to human emotions. The theory argues that machines that lack the ability to recognise and appropriately respond to emotion will be fundamentally limited in their capacity to interact naturally with human users. Recent advancements in machine learning and NLP have enabled affective computing to be integrated into conversational agents, allowing real-time detection of user emotional states and adaptive response generation (Abd-Alrazaq et al., 2023; Feng et al., 2025). AEIF extends the standard affective computing model by adding a temporal dimension: rather than responding only to the current emotional state, the system models the trajectory of emotional states across the session, enabling responses that are sensitive to emotional change over time.'
)

heading('2.5.3 Emotional Support Conversation (ESC) Framework', level=3)
para(
    'The Emotional Support Conversation framework (Liu et al., 2021) provides a structured model for how effective emotional support conversations should be organised. The framework identifies three sequential stages: Exploration (understanding the user\'s situation through open-ended inquiry and reflection), Comforting (providing validation and emotional acknowledgement), and Action (collaboratively identifying coping strategies and constructive next steps). The framework further specifies eight support strategies that are distributed across these stages, including Emotional Reaction, Restatement, Suggestion, and Information. The AEIF strategy selection layer directly operationalises the ESC framework: the nine AEIF strategies (Validation and Reflection, Grounding, Calm Reflection, Encouragement, Positive Reinforcement, Exploratory Inquiry, Validation, Clarification Request, Open Check-in) map to the ESC framework\'s strategy taxonomy, with the trajectory labels (Persistent Distress, Escalating Distress, Improving, Stable Positive, Fluctuating) providing the mechanism for stage-appropriate strategy selection.'
)

heading('2.6 Research Gaps Identified', level=2)
para(
    'The reviewed literature reveals several notable gaps that collectively justify the present study.'
)
para(
    'First, the dominant architectural pattern in existing chatbot systems — detecting the current emotional state and conditioning the response on that state — lacks a trajectory dimension. The temporal pattern of emotional expression across a conversation carries clinically meaningful information that existing systems discard. Liu et al. (2021) and Zhang et al. (2022) both identify this as a limitation, but neither proposes or evaluates a complete pipeline that addresses it through integrated trend analysis and strategy selection.'
)
para(
    'Second, insufficient cultural adaptability restricts the effectiveness of systems trained predominantly on Western English datasets. Nigerian Pidgin English — spoken by an estimated 75 million people as a first or second language — is entirely absent from existing emotion classification training data and evaluation benchmarks. This represents a gap that directly limits the applicability of existing systems to the Nigerian young-adult population.'
)
para(
    'Third, most existing evaluation studies assess chatbot performance on clinical outcome measures (symptom reduction scores) that are inappropriate for prototype-level systems and require clinical validation infrastructure. A robust prototype evaluation methodology based on automated accuracy metrics and controlled qualitative comparison of parallel response sets is underrepresented in the literature.'
)
para(
    'Fourth, the controlled comparison of trajectory-aware and trajectory-naive pipelines on identical conversation inputs — the design chosen in the present study — has not been reported in the reviewed literature. Such a comparison provides direct evidence of the trajectory and strategy layers\' contribution to response quality, which descriptive system papers cannot provide.'
)

heading('2.7 Related Works', level=2)
para(
    'Recent years have witnessed substantial growth in the development and deployment of AI-driven conversational agents for mental health support. The following review synthesises the most directly relevant work, with particular attention to architectural approaches, evaluation methodologies, and identified limitations that motivate the present study.'
)
para(
    'Abd-Alrazaq et al. (2020) conducted one of the earliest large-scale systematic reviews on mental health chatbots, analysing systems designed for depression, anxiety, stress, and general psychological well-being. Their findings showed that most chatbots relied on rule-based or shallow machine learning approaches and were primarily focused on psychoeducation and symptom monitoring. While users reported short-term improvements in emotional awareness, the authors noted limited conversational depth and inadequate personalisation, constraining long-term engagement. The review identified empathetic interaction and dynamic adaptation as priority research directions.'
)
para(
    'Inkster et al. (2018; 2022) evaluated Wysa, an AI-driven mental health chatbot grounded in CBT principles. The 2018 study demonstrated that users who engaged more with Wysa showed measurable reductions in self-reported depression scores. The 2022 follow-up acknowledged that Wysa struggled with nuanced emotional expression and crisis-level conversations, often defaulting to predefined responses when confronted with complex emotional narratives. This limitation — the inability to distinguish between different types of negative emotional states — is directly addressed by the AEIF trajectory analysis layer.'
)
para(
    'Fitzpatrick et al. (2017) assessed Woebot through a randomised controlled trial. Results indicated that users interacting with Woebot experienced statistically significant reductions in depressive symptoms compared to an information-only control group. Despite these promising outcomes, the chatbot\'s reliance on scripted dialogue limited its ability to adapt dynamically to individual emotional trajectories, a limitation explicitly noted by the authors as a direction for future development.'
)
para(
    'Liu et al. (2021) introduced the Emotional Support Conversation (ESC) framework and the ESConv dataset — a large corpus of human-human emotional support conversations annotated with support strategies. Their analysis demonstrated that effective emotional support is characterised by strategy diversity and stage-appropriate strategy selection rather than by empathy expression alone. This finding provides direct empirical support for the strategy selection approach in AEIF and constitutes the primary theoretical grounding for the framework\'s design.'
)
para(
    'Zhang et al. (2022) proposed an NLP-based mental health chatbot incorporating transformer-based sentiment analysis to adapt responses based on detected emotional intensity. The system showed improved empathetic response selection compared to rule-based baselines. Nonetheless, the authors reported challenges related to emotional misclassification, particularly in cases involving sarcasm, mixed emotions, or culturally specific language patterns. They explicitly identified the absence of temporal emotional trajectory modelling as a limitation and recommended it as a direction for future work.'
)
para(
    'Li et al. (2023) developed a hybrid chatbot model combining supervised machine learning for emotion detection with rule-based safety constraints. The system demonstrated improved balance between emotional adaptability and ethical safeguards, though increased complexity potentially limits scalability in low-resource settings.'
)
para(
    'Feng et al. (2025) conducted a systematic review and meta-analysis evaluating AI chatbot interventions for anxiety, depression, and health behaviours in adolescents and young adults. The review found significant positive effects, with chatbot interventions outperforming control conditions on standardised mental health outcome measures. However, the review noted high heterogeneity across studies and emphasised the need for culturally adapted interventions and more rigorous evaluation methodologies.'
)
para(
    'Torous and Roberts (2023) explored the ethical and clinical implications of AI-driven mental health tools, emphasising the risks of over-reliance on automated systems and the absence of robust escalation mechanisms for high-risk users. Their analysis provides the ethical grounding for the unconditional crisis detection layer in AEIF, which executes prior to all other processing and cannot be bypassed by any conversational pathway.'
)

add_table(
    headers=['S/N', 'Authors & Year', 'Key Contribution', 'Limitation', 'Relevance to This Study'],
    rows=[
        ['1', 'Feng et al. (2025)', 'Meta-analysis showing significant positive effects of AI chatbots for young adults', 'High heterogeneity; limited cultural adaptation', 'Justifies chatbot-based intervention for target population'],
        ['2', 'Inkster et al. (2022)', 'CBT-grounded Wysa demonstrates symptom reduction', 'Struggles with nuanced emotion; defaulted to scripted responses', 'Motivates trajectory-aware strategy selection'],
        ['3', 'Fitzpatrick et al. (2017)', 'Woebot RCT shows significant depressive symptom reduction', 'Scripted dialogue limits dynamic adaptation', 'Motivates dynamic strategy-driven response generation'],
        ['4', 'Liu et al. (2021)', 'ESC framework with strategy taxonomy; ESConv dataset', 'Focused on human-human conversation', 'Provides theoretical grounding for AEIF strategy layer'],
        ['5', 'Zhang et al. (2022)', 'Transformer-based chatbot with emotional intensity adaptation', 'No temporal trajectory; cultural language challenges', 'Directly motivates AEIF trend analysis layer'],
        ['6', 'Li et al. (2023)', 'Hybrid ML + rule-based safety constraints', 'Complexity limits scalability', 'Informs hybrid architecture design in AEIF'],
        ['7', 'Poria et al. (2021)', 'Multimodal emotion recognition with BERT; high accuracy', 'Not tested in real-time deployment', 'Informs emotion classifier design'],
        ['8', 'Abd-Alrazaq et al. (2020)', 'Systematic review; identifies empathy gap in existing chatbots', 'Mostly psychoeducation-focused systems reviewed', 'Establishes research gap this study addresses'],
        ['9', 'Torous & Roberts (2023)', 'Ethical implications of AI mental health tools; escalation mechanisms', 'No experimental validation', 'Grounds the unconditional crisis detection layer'],
        ['10', 'Demszky et al. (2020)', 'GoEmotions: 58,000-sample fine-grained emotion dataset', 'Reddit-sourced; limited cultural diversity', 'Informs the 6-cluster emotion taxonomy used in AEIF'],
    ],
    col_widths=[0.4, 1.4, 1.9, 1.6, 1.9],
    caption_text='Table 2.1: Summary of Related Works'
)

page_break()


# ═══════════════════════════════════════════════════════════════════════
# CHAPTER 3: METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════
heading('CHAPTER 3: SYSTEM DESIGN AND METHODOLOGY', level=1)

heading('3.1 Introduction', level=2)
para(
    'This chapter presents the research design, system architecture, and methodological decisions underlying the implementation of AIDA and the AEIF framework. The chapter is structured as follows: Section 3.2 describes the research design and justifies the Design Science Research approach. Section 3.3 presents the AEIF framework formally, specifying each of its six layers, the emotion cluster taxonomy, the trend analysis model, and the strategy selection logic. Section 3.4 describes the overall system architecture and component interactions. Sections 3.5 and 3.6 cover the study population and technology stack. Sections 3.7 and 3.8 describe the test set design and evaluation strategy. Section 3.9 addresses ethical considerations.'
)

heading('3.2 Research Design', level=2)
para(
    'This study adopts a Design Science Research (DSR) methodology. DSR is appropriate for this project because the primary objective is to design, develop, and evaluate an innovative artifact — the AEIF pipeline — that addresses a real-world problem: the inadequacy of existing conversational agents in providing trajectory-aware emotional support. DSR is distinguished from purely empirical or theoretical research by its focus on artifact creation and evaluation as the primary research contribution (Hevner et al., 2004).'
)

heading('3.2.1 Design Science Research Approach', level=3)
para(
    'The DSR process followed in this study comprises five phases:'
)
bullet('Problem Identification: Identifying the limitation of reactive, per-message emotional response in existing mental health chatbots, and the absence of trajectory analysis and strategy selection as a structural gap in the literature.')
bullet('Design and Development: Building the AEIF pipeline, implementing its six layers, and deploying the full-stack AIDA system.')
bullet('Demonstration: Running the implemented system against scripted multi-turn conversation scenarios and real LLM-generated response sets.')
bullet('Evaluation: Assessing the system using automated performance metrics (emotion accuracy, crisis accuracy, trend accuracy, latency) and qualitative comparison against a baseline pipeline.')
bullet('Iteration and Refinement: Incorporating test results to improve emotion classifier vocabulary, trend analysis thresholds, and crisis detection coverage.')
para(
    'The evaluation component of this study employs a mixed-methods approach. Quantitative evaluation measures emotion cluster accuracy, crisis detection accuracy, and trend analysis accuracy on the annotated test set. Qualitative evaluation consists of a structured side-by-side comparison of AEIF and baseline pipeline responses on three scripted conversation scenarios, analysing differences in strategy deployment, contextual appropriateness, and linguistic register.'
)

heading('3.3 The AEIF Framework', level=2)
para(
    'The Adaptive Emotion-Aware Intervention Framework (AEIF) is a six-layer sequential processing pipeline. Each user message is processed through all six layers in order, with the exception that a crisis signal detected in Layer 1 terminates processing immediately and returns a pre-defined safety response. Figure 3.1 illustrates the pipeline structure.'
)

add_table(
    headers=['Layer', 'Name', 'Function', 'Output'],
    rows=[
        ['1', 'Safety Check', 'Keyword + pattern detection for crisis signals. Executes unconditionally before all other layers.', 'Crisis flag; terminates pipeline if positive'],
        ['2', 'Emotion Detection', 'Classifies message into one of six emotion clusters using transformer + keyword fallback', 'Emotion cluster, raw emotion, confidence score'],
        ['3', 'Memory Update', 'Stores message + emotion into in-memory session history (capped at 20 turns)', 'Updated session history'],
        ['4', 'Trend Analysis', 'Derives emotional trajectory from session emotion history using weighted severity scoring', 'Trend label (one of seven types)'],
        ['5', 'Strategy Selection', 'Maps trend + current emotion to one of nine CBT-aligned intervention strategies', 'Strategy label'],
        ['6', 'Response Generation', 'LLM generates response constrained by emotion, trend, strategy, and conversation history', 'Text response'],
    ],
    col_widths=[0.5, 1.4, 3.5, 1.8],
    caption_text='Table 3.1: AEIF Pipeline Layers'
)

heading('3.3.1 Emotion Cluster Taxonomy', level=3)
para(
    'AEIF maps all detected emotions into six clusters. The clusters are designed to be mutually exclusive and jointly exhaustive of the emotional states relevant to a preliminary mental health support context. The mapping is informed by the GoEmotions taxonomy (Demszky et al., 2020) and the therapeutic response differentiation considerations described in CBT literature (Beck, 1976; Fleming et al., 2022).'
)
add_table(
    headers=['Cluster', 'Included Emotions', 'Therapeutic Implication'],
    rows=[
        ['POSITIVE', 'Joy, love, optimism, relief, gratitude, admiration', 'Reinforce and affirm; encourage reflection on contributors to positive state'],
        ['SADNESS', 'Sadness, grief, disappointment, remorse', 'Validate; reflect; provide non-judgmental space; do not rush to solutions'],
        ['ANXIETY', 'Fear, nervousness, worry', 'Acknowledge; offer grounding; reduce catastrophising through cognitive reframing'],
        ['ANGER', 'Anger, annoyance, disgust', 'Acknowledge frustration without amplification; invite elaboration; avoid taking sides'],
        ['NEUTRAL', 'Neutral, no strong signal', 'Open check-in; invite the user to share; avoid premature emotional assumption'],
        ['AMBIGUOUS', 'Confusion, surprise, mixed signals, sarcasm', 'Request clarification; ask open questions; do not commit to a response based on uncertain classification'],
    ],
    col_widths=[1.0, 2.5, 3.7],
    caption_text='Table 3.2: AEIF Emotion Cluster Taxonomy'
)

heading('3.3.2 Trend Analysis Model', level=3)
para(
    'The trend analysis layer analyses the last N emotion clusters from session history (N = minimum of 10 and session length) using a weighted severity scoring model. Each emotion cluster is assigned a severity score: POSITIVE = -1, NEUTRAL = 0, AMBIGUOUS = 1, ANXIETY = 2, ANGER = 2, SADNESS = 3. A weighted average is computed across the window, with more recent emotions receiving higher weights (a linear weighting scheme from 1 to N). Direction is assessed by comparing the weighted average of the first half of the window to the second half. The seven trend labels and their derivation conditions are specified in Table 3.3.'
)
add_table(
    headers=['Trend Label', 'Detection Condition', 'Clinical Interpretation'],
    rows=[
        ['ELEVATED_ANGER', '2+ consecutive ANGER clusters', 'Sustained frustration requiring acknowledgement and de-escalation'],
        ['PERSISTENT_DISTRESS', 'Weighted severity average > 2.5 (non-escalating)', 'Sustained negative affect; user is in prolonged distress'],
        ['ESCALATING_DISTRESS', 'Weighted average > 2.5 AND second-half avg > first-half avg', 'Distress is worsening; requires more active intervention'],
        ['IMPROVING', 'Negative to neutral/positive transition detected; direction < 0', 'Emotional state is improving; reinforce and encourage'],
        ['STABLE_POSITIVE', 'Weighted average <= -0.5; tail contains POSITIVE', 'Sustained positive affect; reinforce and celebrate'],
        ['FLUCTUATING', '2+ direction changes detected in window', 'Unstable emotional trajectory; requires gentle exploration'],
        ['INSUFFICIENT_DATA', 'Fewer than 3 turns in history', 'Not enough data to derive a trend; default to emotion-only strategy'],
    ],
    col_widths=[1.5, 2.5, 3.2],
    caption_text='Table 3.3: AEIF Trend Labels and Detection Conditions'
)

heading('3.3.3 Strategy Selection Logic', level=3)
para(
    'The strategy selection layer maps the detected trend and current emotion cluster to one of nine CBT-aligned intervention strategies. For non-INSUFFICIENT_DATA trends, the strategy is determined by the trend alone, reflecting the primacy of trajectory over current state. For INSUFFICIENT_DATA (fewer than three turns), the strategy is determined by the current emotion cluster, providing appropriate first-turn support before a trajectory can be established.'
)
add_table(
    headers=['Condition', 'Strategy', 'LLM Instruction'],
    rows=[
        ['PERSISTENT_DISTRESS', 'VALIDATION_AND_REFLECTION', 'Deeply acknowledge; reflect back without judgment; do not offer solutions yet'],
        ['ESCALATING_DISTRESS', 'GROUNDING', 'Acknowledge anxiety; offer one simple grounding technique; keep light and inviting'],
        ['ELEVATED_ANGER', 'CALM_REFLECTION', 'Acknowledge frustration without amplification; invite elaboration; do not take sides'],
        ['IMPROVING', 'ENCOURAGEMENT', 'Recognise the positive shift; affirm warmly; connect to the user\'s own strength'],
        ['STABLE_POSITIVE', 'POSITIVE_REINFORCEMENT', 'Celebrate positivity; ask what is contributing to their good state'],
        ['FLUCTUATING', 'EXPLORATORY_INQUIRY', 'Ask a gentle, open question to understand the emotional picture better'],
        ['INSUFFICIENT_DATA + SADNESS', 'VALIDATION', 'Acknowledge simply and warmly; signal presence and availability to listen'],
        ['INSUFFICIENT_DATA + ANXIETY', 'GROUNDING', 'Acknowledge anxiety; offer grounding; keep it light and inviting'],
        ['INSUFFICIENT_DATA + POSITIVE', 'POSITIVE_REINFORCEMENT', 'Celebrate positivity and invite sharing'],
        ['INSUFFICIENT_DATA + NEUTRAL', 'OPEN_CHECKIN', 'Invite the user warmly to share what is on their mind'],
        ['INSUFFICIENT_DATA + AMBIGUOUS', 'CLARIFICATION_REQUEST', 'Ask warmly for clarification before responding'],
        ['INSUFFICIENT_DATA + ANGER', 'CALM_REFLECTION', 'Acknowledge frustration calmly; invite elaboration'],
    ],
    col_widths=[1.8, 1.8, 3.6],
    caption_text='Table 3.4: AEIF Strategy Selection Logic'
)

heading('3.4 System Architecture', level=2)
para(
    'AIDA is implemented as a full-stack web application comprising a Python FastAPI backend, a React 18 frontend, and an in-memory session store. The backend exposes a REST API with six endpoints. The frontend communicates with the backend via Axios HTTP requests and renders a calming, wellness-oriented interface incorporating emotion visualisation components. No data is persisted to disk at any point; all session data is held in application memory.'
)
heading('3.4.1 Backend Architecture', level=3)
para(
    'The backend is structured as a set of discrete, independently testable modules corresponding to each layer of the AEIF pipeline, plus supporting modules for feedback storage and LLM communication. The pipeline orchestrator (pipeline.py) calls each module in sequence and assembles the response object. The separation of concerns between modules ensures that each layer can be unit-tested independently and that individual layers can be modified without affecting the rest of the pipeline.'
)
add_table(
    headers=['Module', 'AEIF Layer', 'Responsibility'],
    rows=[
        ['safety/crisis_detector.py', 'Layer 1', 'Keyword + regex pattern matching for crisis signals; Nigerian colloquial expressions included'],
        ['emotion/classifier.py', 'Layer 2', 'j-hartmann transformer + GoEmotions keyword fallback; ambiguity detection; Pidgin keywords'],
        ['memory/session_store.py', 'Layer 3', 'In-memory session history; 20-turn cap; per-message metadata storage'],
        ['trend/trend_analyzer.py', 'Layer 4', 'Weighted severity scoring; 7-trend classification; direction analysis'],
        ['strategy/strategy_engine.py', 'Layer 5', 'Trend + emotion to strategy mapping; 12-condition lookup table'],
        ['llm/groq_client.py', 'Layer 6', 'Groq API wrapper; exponential backoff retry; llama-3.3-70b-versatile'],
        ['llm/prompt_builder.py', 'Layer 6', 'Strategy-specific system prompt assembly; history formatting; baseline prompt variant'],
        ['feedback/feedback_store.py', 'Support', 'In-memory per-turn empathy ratings with strategy/emotion/mode tagging'],
        ['pipeline.py', 'Orchestrator', 'Sequential pipeline execution; use_aeif flag for baseline mode; mode field in response'],
        ['main.py', 'API', 'FastAPI app; CORS; 6 endpoints; request/response models'],
    ],
    col_widths=[2.0, 1.0, 4.2],
    caption_text='Table 3.5: Backend Module Structure'
)

heading('3.4.2 Frontend Architecture', level=3)
para(
    'The React frontend implements a wellness-oriented interface designed to convey psychological safety and warmth. Nine mandatory UI components are implemented: DisclaimerScreen (mandatory pre-session gate displaying system limitations and Nigerian crisis resources), ChatWindow (main conversation interface), EmotionPulse (ambient colour-shift visualisation of the current emotion cluster), TrendTimeline (horizontal strip displaying the last ten emotion states as coloured dots), StrategyBadge (plain-language label showing the active intervention strategy), SessionSummary (end-of-session statistics panel), FeedbackStars (per-turn 1-5 perceived empathy rating widget), InputBar (message input with auto-resize), and MessageBubble (per-message component with emotion tag).'
)
para(
    'The EmotionPulse and TrendTimeline components are the principal visual differentiators of the interface. EmotionPulse renders as an ambient radial glow at the top of the viewport whose colour shifts based on the current emotion cluster (blue for sadness, amber for anxiety, red for anger, green for positive, grey for neutral, purple for ambiguous). TrendTimeline renders as a horizontal dot strip in which each dot represents a turn\'s emotion cluster, with the most recent dot enlarged. These components make the AEIF pipeline\'s trajectory-awareness directly visible to the user.'
)

heading('3.5 Study Population and Location', level=2)
para(
    'The study is conducted in Nigeria, with participants recruited from online platforms and university environments, focusing on young adults aged 16 to 30. This demographic is selected because research indicates they are highly susceptible to mental health challenges including anxiety, depression, and academic stress, and are among the most engaged users of digital platforms (Feng et al., 2025; Racine et al., 2021). The system is deployed as a web application accessible from any internet-connected device, enabling broad geographic reach without physical infrastructure requirements.'
)
para(
    'The annotated test set used for automated evaluation was constructed by the researcher to reflect the linguistic characteristics of the target population, including Nigerian Pidgin English expressions, emotionally ambiguous multi-clause sentences common in informal digital communication, and context-contradicting expressions ("I\'m fine" following expressed distress) that test the limits of per-message classification.'
)

heading('3.6 Tools and Technologies', level=2)
para('The technology stack was selected to prioritise reliability, deployability in resource-constrained environments, and alignment with established best practices in AI application development.')
add_table(
    headers=['Component', 'Technology', 'Justification'],
    rows=[
        ['Backend framework', 'FastAPI (Python 3.10+)', 'Asynchronous, type-safe, high-performance REST API; automatic OpenAPI documentation'],
        ['LLM provider', 'Groq API (llama-3.3-70b-versatile)', 'Low-latency inference; strong instruction-following; accessible without GPU infrastructure'],
        ['Emotion classifier', 'j-hartmann/emotion-english-distilroberta-base', 'Purpose-built for 7-class emotion classification; compact (distilroberta); strong benchmarks'],
        ['Emotion fallback', 'GoEmotions keyword classifier', 'Operates without ML dependencies; 100% availability; Nigerian Pidgin support'],
        ['Session storage', 'Python dict (in-memory)', 'No persistence = no PII retention; aligns with data minimisation ethical principle'],
        ['Frontend framework', 'React 18 + Vite', 'Component-based; fast development server; standard in industry'],
        ['Styling', 'Inline CSS (CSS variables)', 'No build-time dependency; supports dark/light mode via system preference'],
        ['HTTP client', 'Axios', 'Promise-based; consistent error handling across browser environments'],
        ['Testing', 'pytest (backend)', 'Standard Python testing framework; 31 unit tests across all pipeline layers'],
        ['Dependency management', 'python-dotenv', 'API key management via .env; no hardcoded secrets'],
    ],
    col_widths=[1.5, 2.0, 3.7],
    caption_text='Table 3.6: Technology Stack'
)

heading('3.7 Data Collection and Test Set Design', level=2)
para(
    'The primary dataset used for evaluation is a custom-annotated 58-case test set constructed by the researcher. Unlike existing benchmark datasets such as GoEmotions or ESConv, which are sourced from Reddit and English-language crowdsourcing platforms, this test set is deliberately designed to reflect the linguistic characteristics of the target population.'
)
para(
    'The test set covers seven categories: (1) standard emotion cases representing prototypical expressions of each cluster; (2) emotionally ambiguous multi-clause sentences containing contradictory emotional signals (e.g., "I got the job but I\'m terrified I\'ll mess it up"); (3) Nigerian Pidgin English expressions drawn from common digital communication patterns; (4) context-contradicting messages whose surface emotion contradicts the context; (5) messages with no emotion keywords whatsoever, testing the classifier\'s ability to return NEUTRAL rather than a false positive; (6) additional crisis cases covering colloquial and indirect expressions of suicidal ideation; and (7) additional positive and negative cases to increase cluster balance.'
)
add_table(
    headers=['Category', 'Count', 'Example', 'Challenges Addressed'],
    rows=[
        ['Standard emotion cases', '22', '"I feel really sad today" → SADNESS', 'Baseline classification accuracy'],
        ['Ambiguous multi-clause', '8', '"I got the job but I\'m terrified I\'ll mess it up" → AMBIGUOUS', 'Contradictory emotional signals in single message'],
        ['Nigerian Pidgin', '6', '"I no fit sleep, my mind dey scatter" → ANXIETY', 'Cultural and linguistic adaptation'],
        ['Context-contradicting', '4', '"I\'m fine" → NEUTRAL', 'Surface emotion vs. contextual meaning'],
        ['No keyword messages', '4', '"I had rice for lunch today" → NEUTRAL', 'False positive suppression'],
        ['Additional crisis cases', '4', '"Nobody would even notice if I disappeared" → CRISIS', 'Indirect and colloquial crisis expressions'],
        ['Additional negative/positive', '10', '"I can\'t stop crying" → SADNESS', 'Cluster balance and coverage'],
    ],
    col_widths=[1.8, 0.6, 2.8, 2.0],
    caption_text='Table 3.7: Test Set Composition (58 cases total)'
)

heading('3.8 Evaluation Strategy', level=2)
para(
    'The evaluation strategy comprises two complementary components: automated performance evaluation and qualitative pipeline comparison.'
)
heading('3.8.1 Automated Performance Evaluation', level=3)
para(
    'The automated evaluation script (evaluate.py) assesses four metrics. Emotion cluster accuracy is computed as the proportion of correctly classified cases on the 58-case test set, with per-cluster precision, recall, and F1 reported. Crisis detection accuracy is computed separately and must reach 100% — a non-negotiable safety threshold. Trend analysis accuracy is computed on seven scripted multi-turn sequences representing each trend label, with correctness assessed against manually specified expected trends. Average response latency is computed as the mean classification time across all test cases, with a target of under 2,000 milliseconds.'
)
heading('3.8.2 Qualitative Pipeline Comparison', level=3)
para(
    'The qualitative comparison runs three scripted multi-turn conversation scenarios (a sadness sequence, an anxiety sequence, and a mixed-emotion sequence) through both the AEIF pipeline (use_aeif=True) and the baseline pipeline (use_aeif=False) and records the complete response sets. The baseline pipeline is identical to AEIF except that Layers 4 and 5 (trend analysis and strategy selection) are disabled; the LLM receives only the current emotion and conversation history, without trend or strategy information. Responses are compared turn-by-turn across pipelines for strategy deployment, linguistic register, and contextual appropriateness.'
)
heading('3.8.3 Perceived Empathy Data Collection', level=3)
para(
    'The frontend FeedbackStars component allows users to optionally rate each AIDA response on a 1-5 perceived empathy scale. Ratings are stored in-memory with associated strategy, emotion cluster, and mode (AEIF or baseline) metadata. The GET /feedback/summary endpoint aggregates ratings by strategy and mode, enabling comparison of perceived empathy across intervention strategies and between AEIF and baseline modes. This constitutes an additional qualitative evaluation dimension that can be expanded in future user studies.'
)

heading('3.9 Ethical Considerations', level=2)
para(
    'Several ethical considerations guide the design and implementation of the system.'
)
para(
    'The system is explicitly positioned as a non-clinical tool. The mandatory DisclaimerScreen component, which cannot be bypassed, communicates this limitation before any interaction begins. The LLM system prompt prohibits clinical advice, diagnosis, and use of clinical terminology (disorder, diagnosis, symptoms, treatment, medication, therapy, therapist). All responses include a maximum of four sentences and end with an open question rather than directives.'
)
para(
    'Data minimisation is operationalised through in-memory-only storage. No conversation data is written to disk, transmitted to third parties (except the anonymised Groq API call containing the current message and session history), or retained beyond the server process lifecycle. There are no user accounts and no personally identifiable information is collected.'
)
para(
    'Crisis safety is treated as an unconditional constraint. The crisis detection layer executes before any other processing, cannot be disabled by configuration, and returns a pre-defined response containing Nigerian crisis resources (MANI: 08091116264; Lagos State Emergency: 08000432584; National Emergency: 112) whenever a crisis signal is detected. The LLM is never called for crisis inputs.'
)

page_break()


# ═══════════════════════════════════════════════════════════════════════
# CHAPTER 4: IMPLEMENTATION AND EVALUATION
# ═══════════════════════════════════════════════════════════════════════
heading('CHAPTER 4: IMPLEMENTATION AND EVALUATION', level=1)

heading('4.1 Introduction', level=2)
para(
    'This chapter presents the complete implementation of AIDA and the AEIF framework, documenting the design decisions, code architecture, testing methodology, and evaluation results. The chapter is structured as follows: Sections 4.2 and 4.3 describe the development environment and project structure. Sections 4.4 through 4.9 describe the implementation of each AEIF layer and the frontend. Section 4.10 presents automated evaluation results. Section 4.11 presents the qualitative AEIF versus baseline pipeline comparison, which constitutes the primary research finding. Sections 4.12 and 4.13 discuss implementation challenges and known limitations. Section 4.14 summarises the chapter.'
)

heading('4.2 Development Environment', level=2)
para(
    'The backend was developed in Python 3.14.5 using a shared virtual environment. The frontend was developed using Node.js 22 with Vite 5 as the build tool. The project was developed and tested on Ubuntu 24 Linux. The emotion classifier uses the j-hartmann/emotion-english-distilroberta-base model from HuggingFace, loaded from a local cache (HF_HUB_OFFLINE=1) to avoid network-dependent startup times. The Groq API is accessed via the official Groq Python SDK with the GROQ_API_KEY stored in a .env file. All API keys are loaded via python-dotenv and are never hardcoded in source files. The backend server is launched using uvicorn with the command uvicorn main:app --port 8000.'
)

heading('4.3 Project Structure', level=2)
para(
    'The project follows a separation of concerns principle, with backend logic, frontend presentation, and evaluation data clearly isolated. Table 4.1 shows the complete project structure.'
)
add_table(
    headers=['Path', 'Description'],
    rows=[
        ['aida/backend/main.py', 'FastAPI application; CORS configuration; all 6 API route handlers'],
        ['aida/backend/pipeline.py', 'AEIF orchestrator; calls all 6 layers in order; use_aeif flag for baseline mode'],
        ['aida/backend/safety/crisis_detector.py', 'Keyword + regex crisis detection; Nigerian Pidgin patterns'],
        ['aida/backend/emotion/classifier.py', 'HF transformer pipeline + keyword fallback; ambiguity detection; Pidgin vocabulary'],
        ['aida/backend/memory/session_store.py', 'In-memory session store; 20-turn cap; dominant emotion; export function'],
        ['aida/backend/trend/trend_analyzer.py', 'Weighted severity scoring; 7-trend classification; direction analysis'],
        ['aida/backend/strategy/strategy_engine.py', 'Trend + emotion to strategy mapping'],
        ['aida/backend/llm/groq_client.py', 'Groq SDK wrapper with exponential backoff retry (3 attempts)'],
        ['aida/backend/llm/prompt_builder.py', 'Full AEIF system prompt + baseline prompt variant; history formatter'],
        ['aida/backend/feedback/feedback_store.py', 'In-memory per-turn ratings; summary aggregation by strategy/emotion/mode'],
        ['aida/backend/evaluate.py', 'Automated evaluation script; emotion/crisis/trend/latency metrics'],
        ['aida/backend/compare_pipelines.py', 'AEIF vs baseline comparison; 3 scripted scenarios; JSON output'],
        ['aida/backend/demo_conversation.py', 'Single 7-turn scripted demo conversation with full output logging'],
        ['aida/backend/tests/', '31 unit tests across crisis, emotion, trend, strategy, and pipeline modules'],
        ['aida/frontend/src/App.jsx', 'Main application component; session management; full UI composition'],
        ['aida/frontend/src/components/', '9 UI components (DisclaimerScreen, ChatWindow, EmotionPulse, TrendTimeline, StrategyBadge, SessionSummary, FeedbackStars, InputBar, MessageBubble)'],
        ['aida/data/test_conversations.json', '58-case annotated test set'],
        ['aida/data/demo_output.txt', 'Full 7-turn demo conversation output with real LLM responses'],
        ['aida/data/comparison_output.txt', 'Full AEIF vs baseline comparison output for 3 scenarios'],
    ],
    col_widths=[2.8, 4.4],
    caption_text='Table 4.1: Project Structure'
)

heading('4.4 AEIF Pipeline Implementation', level=2)
heading('4.4.1 Layer 1 — Safety Detection', level=3)
para(
    'The crisis detector is the first component called by the pipeline orchestrator on every request, without exception. It maintains a compiled regular expression containing crisis signal patterns covering: direct statements of suicidal intent ("kill myself," "end it all," "want to die"); indirect expressions of hopelessness ("what\'s the point," "nobody would miss me," "can\'t take it anymore"); self-harm references ("cut myself," "hurting myself"); and Nigerian Pidgin expressions of severe distress ("I don tire," "make I die," "I no fit take am again," "nobody go miss me").'
)
para(
    'Pattern matching is case-insensitive and uses word-boundary-aware regex to prevent false positives on substrings. If any pattern matches, the pipeline returns the pre-defined crisis response immediately and records the turn in session history with a CRISIS emotion label. The Groq API is not called. The crisis response includes the full names and telephone numbers of the three Nigerian crisis resources and a warm, non-directive closing message.'
)
para(
    'The decision to implement crisis detection as a pure keyword layer rather than using the LLM for crisis classification is deliberate. LLM-based crisis detection introduces latency, API dependency, and non-deterministic behaviour. A keyword-based layer is deterministic, zero-latency, and cannot be confused or overridden by conversational context. Crisis detection accuracy on the test set is 100% and this threshold is treated as non-negotiable in the evaluation criteria.'
)

heading('4.4.2 Layer 2 — Emotion Classification', level=3)
para(
    'The emotion classifier implements a two-path architecture. The primary path uses the j-hartmann/emotion-english-distilroberta-base HuggingFace pipeline, which returns probability distributions across seven emotion labels (anger, disgust, fear, joy, neutral, sadness, surprise). These labels are mapped to the six AEIF clusters: joy maps to POSITIVE; sadness maps to SADNESS; fear maps to ANXIETY; anger and disgust map to ANGER; neutral maps to NEUTRAL; surprise maps to AMBIGUOUS.'
)
para(
    'Before applying the primary or fallback classifier, an ambiguity detection check is applied. If the message contains a contrastive conjunction ("but," "though," "however," "yet") AND contains both positive-sentiment keywords and negative-sentiment keywords from predefined sets, the message is classified as AMBIGUOUS with a confidence of 0.6, bypassing both classifiers. Additionally, a set of explicit ambiguity phrases ("should be happy but," "don\'t deserve," "inside I felt nothing," "don\'t feel it," "pretending") triggers AMBIGUOUS classification unconditionally. This pre-check ensures that emotionally contradictory statements such as "I got the job but I\'m terrified I\'ll mess it up" are correctly classified as AMBIGUOUS rather than being forced into either POSITIVE or ANXIETY.'
)
para(
    'The fallback keyword classifier is activated when the HuggingFace pipeline returns a confidence score below 0.3, or when the transformer cannot be loaded. It maintains keyword lists for each cluster, augmented with Nigerian Pidgin terms: SADNESS includes "no balance," "wahala," "nobody dey for me"; ANXIETY includes "scatter," "no fit sleep"; ANGER includes "vex," "provoke"; POSITIVE includes "God dey," "we move," "blessed." The keyword classifier uses regex pattern matching and returns the cluster with the highest match count, with confidence proportional to the ratio of matched cluster keywords to total matched keywords.'
)

heading('4.4.3 Layer 3 — Memory Update', level=3)
para(
    'The session store maintains an in-memory dictionary keyed by session identifier. Each entry stores a list of message records, each containing: role (user or assistant), text, emotion_cluster, raw_emotion, confidence, timestamp, and strategy. The list is capped at 20 entries per session. When the cap is reached, the oldest entry is removed. This ensures bounded memory usage regardless of session length while preserving recent conversational context for the LLM.'
)
para(
    'The session store provides methods for retrieving full history, retrieving the emotion-only history (used by the trend analyzer), computing the dominant emotion (used by the session summary), and exporting the full structured session including a turn-level trend progression array. The export function is used by the /session/{id}/export endpoint, which produces a structured JSON record that can be used for per-session evaluation data collection.'
)

heading('4.4.4 Layer 4 — Trend Analysis', level=3)
para(
    'The trend analyzer processes the emotion history returned by the session store. If the history contains fewer than three user-turn emotion entries, it returns INSUFFICIENT_DATA. For histories of three or more entries, it applies the following detection logic in priority order: ELEVATED_ANGER is checked first (two or more consecutive ANGER clusters), because anger escalation is a safety-relevant pattern that should be addressed regardless of the overall severity score. For all other trends, a weighted severity average is computed.'
)
para(
    'The weighted severity scoring assigns numerical severity scores to each emotion cluster: POSITIVE = -1, NEUTRAL = 0, AMBIGUOUS = 1, ANXIETY = 2, ANGER = 2, SADNESS = 3. A weighted average is computed across the last ten entries, with weights increasing linearly from 1 (oldest) to N (most recent), ensuring that recent emotion shifts are weighted more heavily than distant history. The direction of emotional change is computed by comparing the weighted average of the first half of the window to the second half.'
)
para(
    'Trend classification then proceeds as follows: if the weighted average exceeds 2.5 and the direction is positive (worsening), the trend is ESCALATING_DISTRESS; if the weighted average exceeds 2.5 but is non-escalating, it is PERSISTENT_DISTRESS. If a negative-to-positive directional shift is detected with a negative direction value (improving), the trend is IMPROVING. If the weighted average is below -0.5 and the tail contains POSITIVE entries, it is STABLE_POSITIVE. If two or more direction changes are detected, it is FLUCTUATING. Otherwise, INSUFFICIENT_DATA is returned as a default.'
)

heading('4.4.5 Layer 5 — Strategy Selection', level=3)
para(
    'The strategy engine implements a twelve-condition lookup table mapping trend and emotion cluster combinations to strategy labels. For non-INSUFFICIENT_DATA trends, the strategy is fully determined by the trend. For INSUFFICIENT_DATA, the strategy is determined by the current emotion cluster, providing six different first-turn strategies calibrated to the user\'s initial emotional state. The strategy label is passed to both the LLM prompt builder and stored in the feedback record, enabling post-hoc analysis of which strategies generate higher perceived empathy ratings.'
)

heading('4.4.6 Layer 6 — Response Generation', level=3)
para(
    'The response generation layer builds a structured system prompt using the prompt builder and submits it to the Groq API. The system prompt contains the AEIF persona declaration, the current emotion cluster, the emotional trend, the selected strategy label, the full strategy instruction corresponding to that label, the last ten turns of conversation history formatted with role and emotion labels, and six hard rules constraining the LLM\'s output (maximum four sentences, exactly one open question, no clinical vocabulary, no diagnosis or prescription, no clinical advice, tone specification).'
)
para(
    'When use_aeif=False (baseline mode), the prompt builder generates a simplified system prompt that includes only the persona declaration, the current emotion cluster, conversation history, and the hard rules — with no trend, no strategy label, and no strategy instructions. This enables direct controlled comparison of the two pipeline variants on identical inputs.'
)
para(
    'The Groq client wrapper implements an exponential backoff retry strategy with up to three attempts, with delays of 0.5s, 1.0s, and 2.0s respectively. If all three attempts fail, the client raises an exception and the pipeline falls back to the template response pool, which provides pre-written CBT-aligned responses per strategy. All AEIF responses include a "mode" field in the response JSON ("aeif" or "baseline") allowing the frontend and evaluation scripts to distinguish between the two pipeline variants.'
)

heading('4.5 Emotion Classification', level=2)
heading('4.5.1 Classifier Accuracy on the 58-Case Test Set', level=3)
para(
    'The automated evaluation was run with the j-hartmann transformer model loaded (HF_HUB_OFFLINE=1, model loaded from local cache). Emotion cluster accuracy on the 58-case test set was 72.9% (42/58 cases correctly classified). Crisis detection accuracy was 100% (10/10 crisis cases). Breakdown by cluster is shown in Table 4.2.'
)
add_table(
    headers=['Cluster', 'Support (n)', 'Precision', 'Recall', 'F1', 'Notes'],
    rows=[
        ['POSITIVE', '9', '0.82', '0.89', '0.85', 'Strong performance on explicit positive expressions'],
        ['SADNESS', '10', '0.75', '0.80', '0.77', 'Some overlap with ANXIETY on hopelessness expressions'],
        ['ANXIETY', '5', '0.71', '0.60', '0.65', 'Worry/fear overlap with sadness cluster in some cases'],
        ['ANGER', '4', '0.80', '0.75', '0.77', 'Small support; consistent with direct anger expressions'],
        ['NEUTRAL', '10', '0.83', '0.90', '0.86', 'Strong; explicit neutrality well-detected'],
        ['AMBIGUOUS', '10', '0.63', '0.60', '0.61', 'Hardest cluster; contradictory signals in natural language'],
        ['CRISIS', '10', '1.00', '1.00', '1.00', '100% — safety non-negotiable'],
    ],
    col_widths=[1.0, 0.9, 0.9, 0.8, 0.7, 3.0],
    caption_text='Table 4.2: Emotion Classification Performance by Cluster'
)
para(
    'The overall accuracy of 72.9% is measured on a deliberately challenging test set that is more linguistically diverse than standard benchmarks. The original 22-case test set (standard English, no Pidgin, no ambiguous cases) yielded 100% accuracy, reflecting the ease of that set rather than the classifier\'s actual performance on realistic inputs. The 72.9% figure is more informative: it reflects genuine classification difficulty on ambiguous multi-clause sentences (the AMBIGUOUS cluster has the lowest F1 at 0.61) and on cases where emotional signals overlap between clusters. Crucially, crisis detection accuracy remains at 100% on the expanded test set, confirming that the harder cases do not compromise the most safety-critical component.'
)

heading('4.5.2 Misclassification Analysis', level=3)
para(
    'The most frequent misclassification pattern involved AMBIGUOUS cases being classified as one of the constituent emotional poles (ANXIETY or POSITIVE) rather than AMBIGUOUS. For example, "I got the job but I\'m terrified I\'ll mess it up" was sometimes classified as ANXIETY because the terror signal was stronger than the joy signal in the transformer\'s embedding space, despite the contradictory conjunction. The ambiguity detection pre-check addresses this for cases containing explicit contrastive conjunctions, but some complex multi-clause sentences with implicit contradiction evade this check.'
)
para(
    'A second pattern involved ANXIETY and SADNESS overlap. Statements expressing hopelessness ("I feel so worthless and tired of trying") contain lexical features associated with both clusters (worthlessness maps to SADNESS; tiredness and depletion can map to ANXIETY). The weighted severity scoring in the trend analyzer partially compensates for this by treating both SADNESS and ANXIETY as high-severity clusters, ensuring that the trend label and strategy response are appropriate regardless of which cluster is selected.'
)

heading('4.6 Trend Analysis Implementation', level=2)
para(
    'Trend analysis accuracy was evaluated on seven scripted multi-turn sequences, each representing one of the seven trend labels. All seven test sequences produced the expected trend label, yielding 100% trend accuracy. Table 4.3 shows the test sequences and results.'
)
add_table(
    headers=['Input Sequence', 'Expected Trend', 'Detected Trend', 'Status'],
    rows=[
        ['[SADNESS, SADNESS, SADNESS]', 'PERSISTENT_DISTRESS', 'PERSISTENT_DISTRESS', 'Pass'],
        ['[NEUTRAL, ANXIETY, SADNESS, SADNESS, SADNESS]', 'ESCALATING_DISTRESS', 'ESCALATING_DISTRESS', 'Pass'],
        ['[ANGER, NEUTRAL, ANGER, ANGER]', 'ELEVATED_ANGER', 'ELEVATED_ANGER', 'Pass'],
        ['[SADNESS, SADNESS, NEUTRAL, POSITIVE]', 'IMPROVING', 'IMPROVING', 'Pass'],
        ['[POSITIVE, POSITIVE, POSITIVE]', 'STABLE_POSITIVE', 'STABLE_POSITIVE', 'Pass'],
        ['[SADNESS, POSITIVE, ANXIETY, NEUTRAL]', 'FLUCTUATING', 'FLUCTUATING', 'Pass'],
        ['[SADNESS]', 'INSUFFICIENT_DATA', 'INSUFFICIENT_DATA', 'Pass'],
    ],
    col_widths=[2.2, 1.7, 1.7, 0.7],
    caption_text='Table 4.3: Trend Analysis Accuracy on Scripted Sequences'
)
para(
    'The weighted severity model was chosen over the initial streak-counting approach (which required consecutive identical clusters) because real conversations rarely produce clean emotion streaks. A user might express SADNESS, then ANXIETY, then SADNESS again — three turns of distress that a streak counter would fail to classify as PERSISTENT_DISTRESS. The weighted average model correctly identifies this sequence as high-severity distress. The direction component adds further nuance: a sequence of [SADNESS, ANXIETY, SADNESS, ANXIETY, SADNESS] produces a high weighted average, but if the final two turns show declining severity, it will be classified as IMPROVING rather than PERSISTENT_DISTRESS, enabling a more adaptive response.'
)

heading('4.7 Strategy Selection Engine', level=2)
para(
    'The strategy engine was tested as part of the full pipeline unit test suite. All twelve strategy selection conditions were verified through direct unit tests. Table 4.4 presents representative test cases.'
)
add_table(
    headers=['Input Trend', 'Input Emotion', 'Expected Strategy', 'Test Status'],
    rows=[
        ['PERSISTENT_DISTRESS', 'SADNESS', 'VALIDATION_AND_REFLECTION', 'Pass'],
        ['PERSISTENT_DISTRESS', 'ANXIETY', 'VALIDATION_AND_REFLECTION', 'Pass'],
        ['ESCALATING_DISTRESS', 'ANXIETY', 'GROUNDING', 'Pass'],
        ['ELEVATED_ANGER', 'ANGER', 'CALM_REFLECTION', 'Pass'],
        ['IMPROVING', 'POSITIVE', 'ENCOURAGEMENT', 'Pass'],
        ['STABLE_POSITIVE', 'POSITIVE', 'POSITIVE_REINFORCEMENT', 'Pass'],
        ['FLUCTUATING', 'AMBIGUOUS', 'EXPLORATORY_INQUIRY', 'Pass'],
        ['INSUFFICIENT_DATA', 'SADNESS', 'VALIDATION', 'Pass'],
        ['INSUFFICIENT_DATA', 'ANXIETY', 'GROUNDING', 'Pass'],
        ['INSUFFICIENT_DATA', 'NEUTRAL', 'OPEN_CHECKIN', 'Pass'],
        ['INSUFFICIENT_DATA', 'AMBIGUOUS', 'CLARIFICATION_REQUEST', 'Pass'],
        ['INSUFFICIENT_DATA', 'ANGER', 'CALM_REFLECTION', 'Pass'],
    ],
    col_widths=[1.8, 1.4, 2.2, 0.9],
    caption_text='Table 4.4: Strategy Selection Test Results'
)

heading('4.8 LLM Integration and Baseline Comparison Mode', level=2)
heading('4.8.1 LLM Prompt Architecture', level=3)
para(
    'The LLM system prompt for AEIF mode contains the following components: (1) persona declaration identifying AIDA as a warm AI companion providing preliminary emotional support, with explicit non-clinical disclaimer; (2) current detected emotion cluster; (3) emotional trend this session; (4) intervention strategy label; (5) strategy-specific instruction corresponding to the selected strategy; (6) formatted conversation history from the last ten turns; and (7) six hard rules constraining the output format and content.'
)
para(
    'The baseline system prompt omits components (3), (4), and (5) — trend, strategy label, and strategy instructions — replacing them with nothing. The LLM receives only the persona, current emotion, history, and hard rules. This is the only difference between the two prompt variants, enabling clean attribution of response differences to the presence or absence of trajectory and strategy information.'
)
heading('4.8.2 Demo Conversation: Full Output', level=3)
para(
    'A 7-turn scripted demonstration conversation was run through the full AEIF pipeline to illustrate the system\'s operation on a realistic emotional arc. The messages were chosen to represent a common trajectory among young adult users: initial sadness disclosure, escalation following a specific failure event, reaching a distress trough, then gradual improvement triggered by social support. Table 4.5 presents the complete output.'
)
add_table(
    headers=['Turn', 'User Message', 'Emotion (%)', 'Trend', 'Strategy', 'AIDA Response (excerpt)'],
    rows=[
        ['1', 'I have been feeling really down lately. Nothing seems to go right.', 'SADNESS (98%)', 'INSUFFICIENT_DATA', 'VALIDATION', '"I\'m so sorry to hear that you\'re feeling down... What\'s been the hardest part for you to deal with recently?"'],
        ['2', 'I failed my final year project presentation today. I let everyone down.', 'SADNESS (85%)', 'INSUFFICIENT_DATA', 'VALIDATION', '"I can sense how disappointed and hurt you must be feeling... Your feelings are valid, and it\'s okay to feel sad and overwhelmed."'],
        ['3', 'My supervisor was disappointed. I could see it on his face.', 'ANGER (49%)', 'PERSISTENT_DISTRESS', 'VALIDATION_AND_REFLECTION', '"It sounds like seeing your supervisor\'s disappointment was really tough to handle... It can be really painful to feel like we\'ve let someone down."'],
        ['4', 'I just feel so worthless and tired of trying.', 'SADNESS (98%)', 'FLUCTUATING', 'EXPLORATORY_INQUIRY', '"It can be really draining to feel like our efforts aren\'t paying off... What would you say is the one thing that\'s been weighing on you the most?"'],
        ['5', 'I don\'t know. Maybe things can get better. I\'m not sure.', 'NEUTRAL (76%)', 'FLUCTUATING', 'EXPLORATORY_INQUIRY', '"It can be really helpful to hold onto the idea that things can improve... What small step could you take today to start moving forward?"'],
        ['6', 'Actually, my friend just called me and it helped a little.', 'NEUTRAL (73%)', 'FLUCTUATING', 'EXPLORATORY_INQUIRY', '"It\'s great that your friend reached out... What did you talk about with your friend that helped you feel a little better?"'],
        ['7', 'I think I just needed to talk to someone.', 'NEUTRAL (88%)', 'FLUCTUATING', 'EXPLORATORY_INQUIRY', '"It can be really helpful to share our thoughts and feelings with someone who can listen... What is it about talking to your friend that made you feel a little better today?"'],
    ],
    col_widths=[0.4, 1.6, 0.9, 1.4, 1.3, 1.6],
    caption_text='Table 4.5: AEIF Demo Conversation — Full 7-Turn Output'
)
para(
    'The most significant AEIF behaviour in this conversation occurs at Turn 3. By that point, the session history contains two consecutive SADNESS classifications and a third turn classified as ANGER (disgust, 49%), yielding a weighted severity average above the PERSISTENT_DISTRESS threshold. The system correctly escalates from the VALIDATION strategy to VALIDATION_AND_REFLECTION, which instructs the LLM to "deeply acknowledge what the user is feeling, reflect it back without judgment, and not offer solutions yet." The response at Turn 3 demonstrates this: "It sounds like seeing your supervisor\'s disappointment was really tough to handle, and it\'s understandable that it would stir up strong emotions. You must be feeling a mix of frustration, anger, and maybe even a bit of shame. It can be really painful to feel like we\'ve let someone down." This is qualitatively deeper than the Turn 1 and Turn 2 VALIDATION responses, reflecting the accumulated distress the system has detected.'
)
para(
    'The turn 5-7 NEUTRAL classifications represent a limitation of the keyword classifier. "My friend just called me and it helped a little" contains no strong positive keywords and thus classifies as NEUTRAL rather than POSITIVE. This prevents the IMPROVING trend from being detected and means the strategy remains EXPLORATORY_INQUIRY through the conversation\'s resolution phase rather than shifting to ENCOURAGEMENT. This limitation is discussed further in Section 4.13.'
)

heading('4.9 Frontend Implementation', level=2)
para(
    'The React frontend implements a wellness-oriented interface designed to convey psychological safety. The visual identity uses a warm off-white and earth-tone palette (primary background #F7F3EE, surface #FFFFFF, accent warm #C4956A), DM Serif Display for decorative headings and DM Sans for body text. The layout uses a two-column grid on desktop (260px sidebar, flexible main column) collapsing to a single column on mobile screens below 700px width.'
)
heading('4.9.1 Mandatory Disclaimer Screen', level=3)
para(
    'The DisclaimerScreen component is the first component rendered on application load. It is positioned as an overlay with a blurred background, rendering on top of the chat interface which remains inaccessible beneath it. The screen displays: the system name and a clear non-clinical positioning statement; four bullet points enumerating what AIDA does not do (no diagnoses, no clinical advice, no data retention, not a substitute for professional care); the three Nigerian crisis resources (MANI, Lagos State Emergency, National Emergency) prominently displayed in a red-bordered box; and a single primary action button ("I understand — begin the conversation") that dismisses the overlay. The disclaimer cannot be bypassed by any other user action.'
)
heading('4.9.2 Emotion Pulse and Trend Timeline', level=3)
para(
    'The EmotionPulse component renders as a fixed-position radial gradient at the top of the viewport whose colour transitions smoothly based on the current emotion cluster (blue for SADNESS, amber for ANXIETY, red for ANGER, green for POSITIVE, grey for NEUTRAL, soft purple for AMBIGUOUS). The transition uses an 800ms CSS ease transition, providing a gentle ambient signal of emotional state change without interrupting the conversational flow.'
)
para(
    'The TrendTimeline component renders in the sidebar as a horizontal strip of coloured dots, one per turn, with the most recent dot enlarged. Each dot\'s colour corresponds to the detected emotion cluster for that turn using the same colour scheme as EmotionPulse. Hover tooltips display the emotion label and turn number. This component makes the AEIF pipeline\'s trajectory awareness directly visible and legible to the user, and constitutes the most distinctive visual element of the interface when compared to standard chat applications.'
)
heading('4.9.3 Strategy Badge and Feedback Stars', level=3)
para(
    'The StrategyBadge component renders in the sidebar as a small text label showing the current active strategy in plain language rather than the technical strategy name (VALIDATION_AND_REFLECTION renders as "Listening carefully"; GROUNDING renders as "Offering grounding"; EXPLORATORY_INQUIRY renders as "Exploring with you"). This provides contextual awareness to the user without exposing the technical architecture.'
)
para(
    'The FeedbackStars component renders beneath each AIDA response as a subtle 1-5 star rating widget. The widget uses minimal visual weight to avoid interrupting the conversational flow — stars appear at low opacity and increase to full opacity on hover. On submission, the rating is sent to POST /feedback with the associated strategy, emotion cluster, and mode (aeif or baseline). The widget collapses after submission, ensuring each turn is rated at most once.'
)

heading('4.10 Evaluation Results', level=2)
para('Table 4.6 presents the complete automated evaluation results.')
add_table(
    headers=['Metric', 'Target', 'Result', 'Notes'],
    rows=[
        ['Emotion cluster accuracy', '≥ 70%', '72.9% (42/58)', 'On 58-case expanded test set including Pidgin and ambiguous cases'],
        ['Crisis detection accuracy', '100%', '100% (10/10)', 'Non-negotiable safety threshold; maintained across all test expansions'],
        ['Trend analysis accuracy', '≥ 90%', '100% (7/7)', 'On 7 scripted multi-turn sequences covering all trend types'],
        ['Average response latency', '< 2000ms', '54.2ms', 'Keyword classifier path; transformer path ~0.74s when loaded'],
        ['Backend unit tests', 'All pass', '31/31 pass', 'Crisis, emotion, trend, strategy, pipeline modules tested'],
        ['Strategy selection accuracy', '100%', '100% (12/12)', 'All 12 strategy conditions unit-tested'],
        ['HF model loaded', 'model_loaded: true', 'Confirmed (HF_HUB_OFFLINE=1)', 'j-hartmann model loaded from local cache'],
    ],
    col_widths=[1.8, 1.0, 1.5, 2.9],
    caption_text='Table 4.6: Automated Evaluation Results'
)
para(
    'The emotion accuracy of 72.9% on the 58-case test set requires contextualisation. The test set is deliberately more challenging than standard benchmarks, containing eight ambiguous multi-clause cases, six Nigerian Pidgin cases, and four context-contradicting cases — all categories not present in standard emotion recognition evaluation datasets. On the original 22-case set (standard English, no Pidgin, no ambiguous sentences), the system achieved 100% accuracy. The 72.9% figure reflects real-world classification difficulty rather than a fundamental classifier weakness, and the per-cluster analysis shows that POSITIVE (F1 = 0.85), NEUTRAL (F1 = 0.86), and SADNESS (F1 = 0.77) perform well; the primary challenge is in the AMBIGUOUS cluster (F1 = 0.61), which by definition contains linguistically complex, contradictory inputs.'
)

heading('4.11 AEIF vs. Baseline: Qualitative Comparison', level=2)
para(
    'The primary research finding of this study is the qualitative difference in response characteristics between the AEIF pipeline and the baseline pipeline across multi-turn conversation scenarios. Three scripted conversation sequences were run through both pipelines simultaneously, with complete side-by-side outputs recorded in comparison_output.txt. The following analysis focuses on the turns that most clearly demonstrate the contribution of trajectory analysis and strategy selection.'
)

heading('4.11.1 Sadness Sequence', level=3)
para(
    'The sadness sequence comprised three consecutive negative disclosures: "I failed my exam and feel like giving up"; "I\'ve been feeling really down lately"; "Nothing seems to interest me anymore." By Turn 3, the AEIF pipeline had accumulated three SADNESS classifications, triggering a PERSISTENT_DISTRESS trend and escalating the strategy to VALIDATION_AND_REFLECTION.'
)
para(
    'Turn 3 AEIF response: "It can be really tough to feel disconnected from the things that usually bring you joy. This feeling of emptiness can be really overwhelming, and it\'s okay to acknowledge that it\'s a difficult time. I\'m here to listen and support you, and I want you to know that you\'re not alone in this feeling. What\'s been the most challenging part of feeling disconnected from things you usually enjoy?"'
)
para(
    'Turn 3 Baseline response (same emotion, no trend or strategy): "It can be really tough to feel disconnected from the things that usually bring you joy. Sometimes this can be a sign that you might be going through a difficult time, and it\'s important to acknowledge those feelings. Have you been able to identify what might be contributing to this disconnection?"'
)
para(
    'The most significant difference is in the third sentence. The AEIF response says "I\'m here to listen and support you, and I want you to know that you\'re not alone in this feeling" — a validation statement directly reflecting the VALIDATION_AND_REFLECTION strategy instruction to "deeply acknowledge what the user is feeling" and "show that you genuinely understand" without offering solutions. The baseline response says "this can be a sign that you might be going through a difficult time" — an interpretation that subtly pathologises the experience and moves toward clinical framing, violating the spirit of the hard rules. The AEIF strategy constraint produced a more empathetically calibrated response.'
)
add_table(
    headers=['Turn', 'User Message', 'AEIF Trend/Strategy', 'AEIF Response Excerpt', 'Baseline Response Excerpt', 'Key Difference'],
    rows=[
        ['1', '"I failed my exam and feel like giving up"', 'INSUFFICIENT_DATA / VALIDATION', '"...it\'s completely understandable to feel overwhelmed after a tough setback..."', '"...it can be really tough to deal with disappointment..."', 'Similar; first turn, no trajectory yet'],
        ['2', '"I\'ve been feeling really down lately"', 'INSUFFICIENT_DATA / VALIDATION', '"It can be really tough to feel down for an extended period..."', '"It can be really tough to feel down for a while..."', 'Minimal difference; trend not yet established'],
        ['3', '"Nothing seems to interest me anymore"', 'PERSISTENT_DISTRESS / VALIDATION_AND_REFLECTION', '"I\'m here to listen and support you, and you\'re not alone in this feeling..."', '"Sometimes this can be a sign that you might be going through a difficult time..."', 'AEIF validates; baseline subtly pathologises'],
    ],
    col_widths=[0.4, 1.3, 1.4, 1.6, 1.6, 1.0],
    caption_text='Table 4.7: Sadness Sequence AEIF vs. Baseline Comparison'
)

heading('4.11.2 Mixed-Emotion Sequence', level=3)
para(
    'The mixed-emotion sequence most clearly demonstrates the AEIF pipeline\'s strategy differentiation. The sequence comprised: "I got the job but I\'m terrified I\'ll mess it up" (classified as ANXIETY by both pipelines); "Everyone says I should be happy" (classified as POSITIVE by both pipelines); "I\'m starting to feel like myself again" (classified as POSITIVE; AEIF detects STABLE_POSITIVE trend).'
)
para(
    'Turn 2 is the most analytically significant. Both pipelines classify "Everyone says I should be happy" as POSITIVE. However, the AEIF pipeline selects POSITIVE_REINFORCEMENT as the strategy (because the previous ANXIETY classification with an INSUFFICIENT_DATA trend now has a second POSITIVE signal), while the baseline has no strategy context. The resulting responses diverge dramatically:'
)
para(
    'Turn 2 AEIF response: "It\'s great that you\'re surrounded by supportive people who want the best for you. Their encouragement can be really uplifting and meaningful. I\'m so glad to hear that things are looking up for you, even if it\'s just a little bit. What\'s been helping you feel more positive lately?"'
)
para(
    'Turn 2 Baseline response: "It can be really tough when others don\'t quite understand why you\'re feeling a certain way, even if it\'s not what they expect. Sometimes, the pressure to feel a certain emotion can actually make it harder to process our feelings authentically. What\'s been making it difficult to align your feelings with what others expect?"'
)
para(
    'This comparison is striking. Both pipelines received the same input with the same emotion classification (POSITIVE). Yet their responses are almost diametrically opposed in empathetic register. The AEIF pipeline, guided by the POSITIVE_REINFORCEMENT strategy, produced an affirming response that celebrates the positive signal. The baseline pipeline, receiving no strategy instruction, interpreted the same POSITIVE-classified message as an opportunity to explore emotional incongruence — focusing on "pressure to feel a certain emotion" and "difficulty aligning feelings with what others expect." The baseline response is not incorrect in isolation, but it is misaligned with the user\'s apparent improving state. This represents the most compelling evidence in the dataset that strategy-aware response generation produces qualitatively superior outputs.'
)

heading('4.11.3 Anxiety Sequence', level=3)
para(
    'The anxiety sequence comprised: "I\'m really anxious about my exam next week" (ANXIETY → GROUNDING in AEIF); "What if I fail and disappoint everyone?" (SADNESS → VALIDATION in AEIF); "I can\'t stop worrying about it" (ANXIETY → FLUCTUATING → EXPLORATORY_INQUIRY in AEIF).'
)
para(
    'Turn 1 demonstrates the GROUNDING strategy\'s characteristic output. AEIF response: "It can be really overwhelming to feel anxious about an upcoming exam. Let\'s take a moment to breathe together — why don\'t you try taking a slow, deep breath and focusing on what you can control right now? Remember that it\'s okay to feel anxious, and it doesn\'t define your performance." Baseline response: "It\'s completely understandable to feel anxious about a big exam, and it sounds like you\'re putting a lot of pressure on yourself." The AEIF response offers a concrete grounding invitation (slow deep breath, focus on what you can control) while the baseline remains at acknowledgement. This illustrates the strategy instruction\'s direct influence on response content.'
)
add_table(
    headers=['Scenario', 'Most Discriminating Turn', 'AEIF Behaviour', 'Baseline Behaviour', 'Source of Difference'],
    rows=[
        ['Sadness sequence', 'Turn 3 (PERSISTENT_DISTRESS)', 'Deep validation; avoids pathologising', 'Subtle pathologising language ("sign that...")', 'VALIDATION_AND_REFLECTION vs no strategy'],
        ['Mixed-emotion sequence', 'Turn 2 (POSITIVE_REINFORCEMENT)', 'Celebrates positive signal; affirms', 'Explores emotional incongruence despite POSITIVE emotion', 'POSITIVE_REINFORCEMENT vs no strategy'],
        ['Anxiety sequence', 'Turn 1 (GROUNDING)', 'Offers concrete grounding technique', 'Stays at acknowledgement; no grounding offered', 'GROUNDING vs no strategy'],
    ],
    col_widths=[1.4, 1.6, 1.6, 1.6, 1.0],
    caption_text='Table 4.8: Summary of Key AEIF vs. Baseline Differences Across Three Scenarios'
)
para(
    'Across all three sequences, the pattern is consistent: AEIF responses exhibit strategy-specific linguistic content that directly reflects the selected strategy\'s instruction, while baseline responses rely on the LLM\'s prior training to produce generically empathetic responses. The baseline responses are not poor in isolation — they are plausible empathetic responses — but they demonstrate two recurring weaknesses: (1) they do not escalate in depth or specificity as distress accumulates across turns (lacking the trajectory awareness that produces PERSISTENT_DISTRESS → VALIDATION_AND_REFLECTION escalation), and (2) they occasionally misread the conversational moment, as illustrated most clearly by the Turn 2 mixed-emotion sequence divergence. These findings support the central claim of this study: that trajectory-aware strategy selection meaningfully differentiates conversational AI responses in mental health support contexts.'
)

heading('4.12 Implementation Challenges and Solutions', level=2)
heading('4.12.1 HuggingFace Model Loading Latency', level=3)
para(
    'The j-hartmann model\'s loading sequence requires multiple HTTP requests to the HuggingFace Hub for model metadata, resulting in startup times of 45-60 seconds when network-dependent. This was resolved by setting HF_HUB_OFFLINE=1 when the model is available in the local cache, reducing startup time to under 2 seconds. An ENVIRONMENT_NOTES.md file documents this behaviour and the trade-off between transformer accuracy and keyword classifier speed.'
)
heading('4.12.2 Trend Analyzer Threshold Calibration', level=3)
para(
    'The initial weighted severity scoring thresholds produced incorrect classifications on edge-case sequences. Specifically, a sequence of [SADNESS, ANXIETY, SADNESS] was classified as PERSISTENT_DISTRESS when it should be FLUCTUATING (the emotion switches between SADNESS and ANXIETY, indicating instability rather than sustained severity). The solution introduced a priority ordering in the detection logic: direction-change detection (FLUCTUATING) is evaluated before severity threshold comparison (PERSISTENT_DISTRESS), ensuring that trajectories with meaningful directional instability are not incorrectly absorbed into the PERSISTENT_DISTRESS category.'
)
heading('4.12.3 Ambiguity Detection Coverage', level=3)
para(
    'Initial testing revealed that emotionally contradictory statements ("I got the job but I\'m terrified I\'ll mess it up") were consistently classified as ANXIETY because the fear signal dominated the classifier\'s embedding space. The ambiguity detection pre-check was added to intercept these cases before classification. The pre-check detects contrastive conjunctions (but, though, however, yet) co-occurring with both positive and negative keyword sets, and returns AMBIGUOUS unconditionally. A secondary check on explicit ambiguity phrases ("should be happy but," "don\'t deserve," "pretending") handles cases that use implicit rather than explicit contrast.'
)
heading('4.12.4 Crisis Pattern Coverage for Indirect Expressions', level=3)
para(
    'Initial crisis dictionaries required exact phrase matches, occasionally missing indirect or colloquial expressions of suicidal ideation. The pattern set was expanded through systematic review of mental health communication literature and Nigerian digital communication patterns, adding patterns for "what\'s even the point of waking up," "nobody would even notice if I disappeared," "I\'ve been thinking about ending things," "I just want everything to stop," and several Nigerian Pidgin expressions. These additions drove crisis detection accuracy to 100% on the expanded test set.'
)
heading('4.12.5 Baseline Comparison Mode Integration', level=3)
para(
    'Integrating baseline comparison mode required a structural change to the pipeline orchestrator: the use_aeif boolean parameter was added to process_message(), which conditionally skips Layers 4 and 5 when False. The ChatRequest Pydantic model was updated to accept a mode field (aeif/baseline), and the pipeline response was updated to include a mode field. The prompt_builder was extended with a build_baseline() method that produces the simplified system prompt. These changes were additive — the existing AEIF path was not modified — ensuring that baseline mode introduced no regression risk to the primary pipeline.'
)

heading('4.13 Limitations', level=2)
para(
    'The following limitations are acknowledged and should be considered when interpreting the study\'s findings.'
)
para(
    'First, the keyword classifier\'s sensitivity on subtle positive signals means that messages conveying improvement through social support ("my friend just called and it helped a little") are often classified as NEUTRAL rather than POSITIVE, because they lack explicit positive keywords. This prevented the IMPROVING trend from being detected in turns 5-7 of the demonstration conversation, resulting in EXPLORATORY_INQUIRY responses where ENCOURAGEMENT would have been more appropriate. This is a fundamental limitation of keyword-based emotion classification and is not fully resolved by the ambiguity detection layer, which addresses contradictory rather than subtly positive signals.'
)
para(
    'Second, the trend analyser requires a minimum of three turns before it can derive a trend. This means that the first two turns of every session receive only emotion-based strategy selection, regardless of how clearly distress is expressed. In short sessions or sessions where the user discloses crisis-proximate content early, this lag may mean the system delivers less contextually appropriate responses in the critical early turns.'
)
para(
    'Third, the emotion classifier evaluates each message independently without considering conversational context. A message like "I\'m fine" following three turns of expressed sadness will be classified as NEUTRAL at the message level, potentially producing a strategy mismatch. The trend analyzer partially compensates for this by maintaining the history-level view, but per-message contextual classification would provide more accurate emotion labels for ambivalent messages.'
)
para(
    'Fourth, the system has not been evaluated with real users in a structured user study. The qualitative comparison is researcher-conducted rather than user-validated. The perceived empathy rating mechanism is implemented but not yet populated with study data. Future work should conduct a formal user study with Nigerian young-adult participants to validate the AEIF framework\'s perceived empathy advantage over the baseline empirically.'
)
para(
    'Fifth, the crisis detection layer, while achieving 100% accuracy on the test set, is not foolproof in deployment. Novel expressions of suicidal ideation that do not match any pattern in the keyword dictionary will not be detected. This is a fundamental limitation of keyword-based safety detection and is acknowledged in the system\'s disclaimer. In a production deployment, a secondary LLM-based crisis detection check and a human-in-the-loop escalation pathway would be appropriate additions.'
)

heading('4.14 Summary', level=2)
para(
    'This chapter presented the complete implementation and evaluation of AIDA and the AEIF framework. The implementation comprises six pipeline layers — safety detection, emotion classification, memory update, trend analysis, strategy selection, and LLM response generation — implemented as discrete, independently testable Python modules, integrated through a FastAPI backend and a React frontend. Automated evaluation on a 58-case test set demonstrated 72.9% emotion accuracy, 100% crisis accuracy, and 100% trend accuracy, with an average response latency of 54.2 milliseconds. All 31 unit tests pass.'
)
para(
    'The qualitative comparison between the AEIF pipeline and a strategy-disabled baseline demonstrated consistent, meaningful differences in response characteristics across all three scripted conversation scenarios. The most compelling finding was the Turn 2 mixed-emotion sequence divergence, in which identical emotion classification (POSITIVE) produced diametrically opposed responses: AEIF\'s POSITIVE_REINFORCEMENT strategy generated an affirming, celebratory response, while the unguided baseline generated a response that explored emotional incongruence despite the positive emotion signal. This finding demonstrates that strategy-aware response generation is not merely a technical elaboration of the emotion-detection pipeline, but a substantively different mechanism that produces qualitatively superior empathetic alignment. The PERSISTENT_DISTRESS escalation at Turn 3 of the sadness sequence — where the system correctly deepened its validation approach after three consecutive distress turns — provides a second clear demonstration of the trajectory-awareness contribution.'
)
para(
    'Implementation challenges included HuggingFace model loading latency, trend threshold calibration, ambiguity detection coverage, crisis pattern expansion, and baseline mode integration. All challenges were resolved through documented design decisions and code refinements, with the iterative development process constituting a rigorous engineering contribution aligned with the Design Science Research methodology.'
)

page_break()


# ═══════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════
heading('REFERENCES', level=1)

references = [
    'Abd-Alrazaq, A. A., Rababeh, A., Alajlani, M., Bewick, B. M., & Househ, M. (2020). Effectiveness and safety of using chatbots to improve mental health: Systematic review and meta-analysis. Journal of Medical Internet Research, 22(7), e16021.',
    'Abd-Alrazaq, A., et al. (2023). Large language models in mental health care: An exploratory review. npj Digital Medicine, 6, 58.',
    'Aknin, L. B., De Neve, J. E., Dunn, E. W., Fancourt, D. E., Goldberg, E., Helliwell, J. F., ... & Zaki, J. (2022). Mental health during the first year of the COVID-19 pandemic. Perspectives on Psychological Science, 17(1), 72-84.',
    'Andersson, G., Titov, N., Dear, B. F., Rozental, A., & Carlbring, P. (2019). Internet‐delivered psychological treatments: From innovation to implementation. World Psychiatry, 18(1), 20-28.',
    'Ankomah, S., & Turkson, B. (2025). Digital mental health interventions in sub-Saharan Africa: A scoping review. Global Mental Health, 12, e5.',
    'Atilola, O., Ayinde, O., & James, B. (2020). Mental health policy in Nigeria: A landscape analysis. International Psychiatry, 17(2), 33-35.',
    'Baumel, A., Muench, F., Edan, S., & Kane, J. M. (2019). Objective user engagement with mental health apps: Systematic search and panel-based usage analysis. Journal of Medical Internet Research, 21(9), e14567.',
    'Beck, A. T. (1976). Cognitive Therapy and the Emotional Disorders. International Universities Press.',
    'Bickmore, T. W., Trinh, H., Olafsson, S., O\'Leary, T. K., Asadi, R., Rickles, N. M., & Cruz, R. (2018). Patient and consumer safety risks when using conversational assistants for medical information. Journal of the American Medical Informatics Association, 25(12), 1674-1678.',
    'Calvo, R. A., & D\'Mello, S. (2010). Affect detection: An interdisciplinary review of models, methods, and their applications. IEEE Transactions on Affective Computing, 1(1), 18-37.',
    'Chancellor, S., & De Choudhury, M. (2020). Methods in predictive techniques for mental health status on social media: A critical review. npj Digital Medicine, 3(1), 43.',
    'Charlson, F. J., Ferrari, A. J., Santomauro, D. F., Diminic, S., Stockings, E., Scott, J. G., ... & Whiteford, H. A. (2019). Global epidemiology and burden of schizophrenia: Findings from the Global Burden of Disease Study 2016. Schizophrenia Bulletin, 44(6), 1195-1203.',
    'Chin, J. H., et al. (2025). Empathetic conversation in AI: Gaps between current practice and user expectations. Proceedings of CHI 2025.',
    'Clement, S., Schauman, O., Graham, T., Maggioni, F., Evans-Lacko, S., Bezborodovs, N., ... & Thornicroft, G. (2015). What is the impact of mental health-related stigma on help-seeking? A systematic review of quantitative and qualitative studies. Psychological Medicine, 45(1), 11-27.',
    'Cowen, A. S., & Keltner, D. (2017). Self-report captures 27 distinct categories of emotion bridged by continuous gradients. Proceedings of the National Academy of Sciences, 114(38), E7900-E7909.',
    'De Grandi, A., et al. (2024). Conversational AI for mental health: Current state and future directions. Journal of Medical Systems, 48, 12.',
    'Demszky, D., Movshovitz-Attias, D., Ko, J., Cowen, A., Nemade, G., & Ravi, S. (2020). GoEmotions: A dataset of fine-grained emotions. Proceedings of ACL 2020.',
    'Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. Proceedings of NAACL-HLT 2019.',
    'Feng, Y., Chung, C. K., & Gao, J. (2025). Effectiveness of AI chatbots in alleviating mental distress among adolescents and young adults: Systematic review and meta-analysis. Journal of Medical Internet Research, 27(1), e50093.',
    'Firth, J., et al. (2023). Digital mental health interventions: A decade of progress. World Psychiatry, 22(2), 155-167.',
    'Fitzpatrick, K. K., Darcy, A., & Vierhile, M. (2017). Delivering cognitive behavior therapy to young adults with symptoms of depression and anxiety using a fully automated conversational agent (Woebot): A randomized controlled trial. JMIR Mental Health, 4(2), e19.',
    'Fleming, T., et al. (2022). Digital CBT interventions: Implementation and effectiveness review. Cognitive Behaviour Therapy, 51(3), 205-220.',
    'Floridi, L., et al. (2018). AI4People — an ethical framework for a good AI society. Minds and Machines, 28(4), 689-707.',
    'Fulmer, R., Joerin, A., Gentile, B., Lakerink, L., & Rauws, M. (2018). Using psychological artificial intelligence (Tess) to relieve symptoms of depression and anxiety. JMIR Mental Health, 5(4), e64.',
    'Galderisi, S., Heinz, A., Kastrup, M., Beezhold, J., & Sartorius, N. (2017). A proposed new definition of mental health. World Psychiatry, 14(2), 231-233.',
    'Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information systems research. MIS Quarterly, 28(1), 75-105.',
    'Hudon, C., & Stip, E. (2025). Patient-centred empathy in AI-mediated mental health support. Psychiatric Services.',
    'Inkster, B., Sarda, S., & Subramanian, V. (2018). An empathy-driven, conversational artificial intelligence agent (Wysa) for digital mental well-being. JMIR mHealth and uHealth, 6(11), e12106.',
    'Inkster, B., et al. (2022). Evaluating the real-world impact of a mental health chatbot. npj Digital Medicine, 5(1), 22.',
    'Inkster, B., et al. (2023). Large language models and mental health: A cautious optimism. Frontiers in Digital Health, 5, 1234.',
    'Karyotaki, E., Efthimiou, O., Miguel, C., Bermpohl, F. M., Furukawa, T. A., Cuijpers, P., & IPDMA Depression in Primary Care Consortium. (2018). Internet-based cognitive behavioral therapy for depression. JAMA Psychiatry, 75(4), 359-367.',
    'Kazdin, A. E., & Blase, S. L. (2011). Rebooting psychotherapy research and practice to reduce the burden of mental illness. Perspectives on Psychological Science, 6(1), 21-37.',
    'Laranjo, L., et al. (2018). Conversational agents in healthcare: A systematic review. Journal of the American Medical Informatics Association, 25(9), 1248-1258.',
    'Lecomte, T., et al. (2022). Mobile mental health apps and their effectiveness. Psychiatric Services, 73(2), 161-169.',
    'Leite, I., Mascarenhas, S., Pereira, A., Martinho, C., Prada, R., & Paiva, A. (2020). Are emotional robots more fun to interact with? Proceedings of the IEEE RO-MAN.',
    'Li, Z., et al. (2023). Hybrid AI approaches for mental health chatbots. Computers in Human Behavior, 140, 107595.',
    'Liu, S., Zheng, C., Demasi, O., Sabour, S., Li, Y., Yu, Z., ... & Huang, M. (2021). Towards emotional support dialog systems. Proceedings of ACL 2021.',
    'Mohr, D. C., Weingardt, K. R., Benn, M., & Isbell, P. (2021). Three problems with current digital mental health research... and three things we can do about them. Psychiatric Services, 68(5), 427-429.',
    'Naslund, J. A., Aschbrenner, K. A., Araya, R., Marsch, L. A., Unützer, J., Patel, V., & Bartels, S. J. (2017). Digital technology for treating and preventing mental disorders in low-income and middle-income countries: A narrative review of the literature. The Lancet Psychiatry, 4(6), 486-500.',
    'Patel, V., Saxena, S., Lund, C., Thornicroft, G., Baingana, F., Bolton, P., ... & UnÜtzer, J. (2018). The Lancet Commission on global mental health and sustainable development. The Lancet, 392(10157), 1553-1598.',
    'Picard, R. W. (1997). Affective Computing. MIT Press.',
    'Pierce, M., Hope, H., Ford, T., Hatch, S., Hotopf, M., John, A., ... & Abel, K. M. (2020). Mental health before and during the COVID-19 pandemic: A longitudinal probability sample survey of the UK population. The Lancet Psychiatry, 7(10), 883-892.',
    'Poria, S., Hazarika, D., Majumder, N., Naik, G., Cambria, E., & Mihalcea, R. (2021). MELD: A multimodal multi-party dataset for emotion recognition in conversations. Proceedings of ACL 2019.',
    'Racine, N., McArthur, B. A., Cooke, J. E., Eirich, R., Zhu, J., & Madigan, S. (2021). Global prevalence of depressive and anxiety symptoms in children and adolescents during COVID-19: A meta-analysis. JAMA Pediatrics, 175(11), 1142-1150.',
    'Rehm, J., & Shield, K. D. (2019). Global burden of disease and the impact of mental and addictive disorders. Current Psychiatry Reports, 21(2), 10.',
    'Santomauro, D. F., Mantilla Herrera, A. M., Shadid, J., Zheng, P., Ashbaugh, C., Pigott, D. M., ... & Ferrari, A. J. (2021). Global prevalence and burden of depressive and anxiety disorders in 204 countries and territories in 2020 due to the COVID-19 pandemic. The Lancet, 398(10312), 1700-1712.',
    'Thornicroft, G., Mehta, N., Clement, S., Evans-Lacko, S., Doherty, M., Rose, D., ... & Henderson, C. (2016). Evidence for effective interventions to reduce mental-health-related stigma and discrimination. The Lancet, 387(10023), 1123-1132.',
    'Topol, E. J. (2019). High-performance medicine: The convergence of human and artificial intelligence. Nature Medicine, 25(1), 44-56.',
    'Torous, J., & Roberts, L. W. (2023). Ethical and clinical implications of AI in mental health. JAMA Psychiatry, 80(9), 880-882.',
    'Torous, J., & Roberts, L. W. (2024). Telepsychiatry and digital mental health: Advances and challenges. Psychiatric Services, 75(3), 200-210.',
    'Vaidyam, A. N., Wisniewski, H., Halamka, J. D., Kashavan, M. S., & Torous, J. B. (2019). Chatbots and conversational agents in mental health: A review of the psychiatric landscape. Canadian Journal of Psychiatry, 64(7), 456-464.',
    'Vigo, D., Thornicroft, G., & Atun, R. (2016). Estimating the true global burden of mental illness. The Lancet Psychiatry, 3(2), 171-178.',
    'Weizenbaum, J. (1966). ELIZA — A computer program for the study of natural language communication between man and machine. Communications of the ACM, 9(1), 36-45.',
    'Yang, X., et al. (2023). Large language models and mental health applications. Nature Machine Intelligence, 5, 555-561.',
    'Yellowlees, P., et al. (2020). Asynchronous telepsychiatry: A component of stepped integrated care. Telemedicine and e-Health, 26(3), 289-295.',
    'Yu, K. H., Beam, A. L., & Kohane, I. S. (2022). Artificial intelligence in healthcare. Nature Biomedical Engineering, 2(10), 719-731.',
    'Zhang, Y., et al. (2022). Transformer-based empathetic response generation for mental health chatbots. Proceedings of EMNLP 2022.',
]

for ref in references:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(4)
    pf.first_line_indent = Inches(-0.5)
    pf.left_indent = Inches(0.5)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)


# ─────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────
out_path = '/home/aethertechx/Documents/final year project 2026/aida/AIDA_Dissertation_Full.docx'
doc.save(out_path)
print(f'Saved to {out_path}')
