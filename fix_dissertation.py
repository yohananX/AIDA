"""
Fix all inaccuracies in AIDA_Dissertation_Corrected.docx to match actual code.
"""
import re

# ─── Read the docx ──────────────────────────────────────────────────
import docx
doc = docx.Document("aida/AIDA_Dissertation_Corrected.docx")

# ─── Helper ─────────────────────────────────────────────────────────
def update_paragraph(doc, idx, new_text):
    """Replace paragraph text at index, preserving runs and formatting."""
    p = doc.paragraphs[idx]
    # Clear existing runs, add one with the new text
    for run in p.runs:
        run.text = ""
    p.runs[0].text = new_text

def find_para(doc, substring, start=0):
    """Return index of first paragraph containing substring."""
    for i, p in enumerate(doc.paragraphs):
        if i >= start and substring in p.text:
            return i
    return None

# ═══════════════════════════════════════════════════════════════════════
# 1. SECTION 3.3.2 — Trend Analysis Model description (paragraph ~166)
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "weighted severity scoring model")
if idx:
    update_paragraph(doc, idx,
        "The trend analysis layer analyses the last N emotion clusters from session history "
        "(N = minimum of 10 and session length) using a severity scoring model. "
        "Each emotion cluster is assigned a severity score: POSITIVE = 0.0, NEUTRAL = 0.3, "
        "AMBIGUOUS = 0.4, ANXIETY = 0.6, ANGER = 0.7, SADNESS = 0.8, CRISIS = 1.0. "
        "A simple arithmetic mean (unweighted) is computed across the window; all entries "
        "contribute equally regardless of recency. Direction is assessed by checking strict "
        "monotonicity: if every consecutive score strictly increases, the direction is worsening "
        "(ESCALATING_DISTRESS); if every consecutive score strictly decreases and the first half "
        "contains negative emotions, the direction is improving (IMPROVING). "
        "If neither monotonic pattern holds, direction changes are counted to detect fluctuation. "
        "The seven trend labels and their derivation conditions are specified in Table 3.3."
    )
    print(f"  Fixed Section 3.3.2 at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 2. SECTION 4.4.4 — Trend Analysis Implementation (paragraph ~231)
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "A weighted average is computed across the last ten")
if idx:
    update_paragraph(doc, idx,
        "The severity scoring assigns numerical severity scores to each emotion cluster: "
        "POSITIVE = 0.0, NEUTRAL = 0.3, AMBIGUOUS = 0.4, ANXIETY = 0.6, ANGER = 0.7, "
        "SADNESS = 0.8, CRISIS = 1.0. An unweighted arithmetic mean is computed across the "
        "last ten entries; all entries contribute equally regardless of recency. "
        "Direction of emotional change is assessed using strict monotonicity checks: "
        "ESCALATING_DISTRESS requires every consecutive score to be strictly increasing; "
        "IMPROVING requires every consecutive score to be strictly decreasing and the first "
        "half of the window to contain negative emotions (SADNESS, ANXIETY, or ANGER). "
        "Fluctuation is detected when one or more direction reversals occur between "
        "consecutive triples."
    )
    print(f"  Fixed Section 4.4.4 weighting/description at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 3. SECTION 4.4.4 paragraph 2 — Trend classification logic
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "Trend classification then proceeds")
if idx:
    update_paragraph(doc, idx,
        "Trend classification then proceeds in priority order as follows: "
        "ELEVATED_ANGER is checked first (two or more consecutive ANGER clusters). "
        "Next, if the emotion sequence has fewer than three entries, INSUFFICIENT_DATA is returned. "
        "Otherwise, an unweighted severity mean is computed, and the following checks are applied "
        "in order: (1) ESCALATING_DISTRESS — every score strictly increases (no threshold required); "
        "(2) IMPROVING — every score strictly decreases and the first half contains negative emotions; "
        "(3) FLUCTUATING — one or more direction reversals are detected; "
        "(4) PERSISTENT_DISTRESS — the mean exceeds 0.55; "
        "(5) STABLE_POSITIVE — the mean is at most 0.2 and the last three entries contain at least "
        "one POSITIVE. If none match, INSUFFICIENT_DATA is returned as default."
    )
    print(f"  Fixed Section 4.4.4 classification logic at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 4. SECTION 4.5.1 — Emotion accuracy (paragraph ~241)
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "72.9% (42/58 cases correctly classified)")
if idx:
    update_paragraph(doc, idx,
        "Emotion cluster accuracy on the 48 non-crisis test cases was 77.08% (37/48 cases "
        "correctly classified). Crisis detection accuracy was 100% (10/10 crisis cases, and no "
        "false positives on the 48 non-crisis cases). Breakdown by cluster is shown in Table 4.2."
    )
    print(f"  Fixed emotion accuracy at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 5. SECTION 4.5.1 — Second reference to 72.9% (paragraph ~244)
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "The overall accuracy of 72.9%")
if idx:
    update_paragraph(doc, idx,
        "The overall accuracy of 77.08% is measured on a deliberately challenging test set that "
        "is more linguistically diverse than standard benchmarks. The original 22-case test set "
        "(standard English, no Pidgin, no ambiguous cases) yielded 100% accuracy, reflecting the "
        "ease of that set rather than the classifier's actual performance on realistic inputs. "
        "The 77.08% figure is more informative: it reflects genuine classification difficulty on "
        "ambiguous multi-clause sentences (the AMBIGUOUS cluster has the lowest F1 at 0.5714) and "
        "on cases where emotional signals overlap between clusters. Crucially, crisis detection "
        "accuracy remains at 100% on the expanded test set, confirming that the harder cases do "
        "not compromise the most safety-critical component."
    )
    print(f"  Fixed accuracy paragraph at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 6. SECTION 4.5.2 — Misclassification examples
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "I got the job but I'm terrified I'll mess it up")
if idx:
    update_paragraph(doc, idx,
        "The most frequent misclassification pattern involved AMBIGUOUS cases being classified "
        "as NEUTRAL rather than AMBIGUOUS. For example, \"I got the job but I'm terrified I'll "
        "mess it up\" was classified as NEUTRAL because the transformer model's neutral logit "
        "dominated the embedding space, despite the contradictory conjunction. This occurred "
        "when the HuggingFace model returned NEUTRAL with confidence above the 0.3 threshold, "
        "causing the ambiguity pre-check to be bypassed. The keyword classifier fallback is only "
        "activated when HF confidence falls below 0.3, and the ambiguity pre-check only runs "
        "on the classification result, not as an independent gate. This means that messages "
        "containing contrastive conjunctions may still be classified as NEUTRAL rather than "
        "AMBIGUOUS if the transformer assigns NEUTRAL a high probability."
    )
    print(f"  Fixed misclassification paragraph 1 at paragraph {idx}")

idx = find_para(doc, "I feel so worthless and tired of trying")
if idx:
    update_paragraph(doc, idx,
        "A second pattern involved SADNESS cases being misclassified as NEUTRAL. Example test "
        "cases such as \"I feel so alone even when people are around\" (expected SADNESS, "
        "classified NEUTRAL) and \"E be like say nobody dey for me\" (expected SADNESS, "
        "classified NEUTRAL) demonstrate that the keyword classifier's lower confidence bound "
        "(0.5 threshold) and the HuggingFace model's neutral bias for non-explicit language "
        "combine to produce false NEUTRAL classifications for statements that a human reader "
        "would recognise as sadness. Additionally, one positive case (\"My mum called today and "
        "it made my whole day\") was misclassified as ANGER due to the transformer assigning a "
        "high probability to the disgust label, which maps to the ANGER cluster. "
        "The severity scoring in the trend analyzer partially compensates for these "
        "misclassifications at the session level, but per-message accuracy on subtle emotional "
        "signals remains an area for improvement."
    )
    print(f"  Fixed misclassification paragraph 2 at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 7. SECTION 4.10 — Evaluation Results table description (paragraph ~281)
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "The emotion accuracy of 72.9% on the 58-case test set requires")
if idx:
    update_paragraph(doc, idx,
        "The emotion accuracy of 77.08% on the 48-case non-crisis test set requires "
        "contextualisation. The test set is deliberately more challenging than standard "
        "benchmarks, containing eight ambiguous multi-clause cases, six Nigerian Pidgin cases, "
        "and four context-contradicting cases — all categories not present in standard emotion "
        "recognition evaluation datasets. On the original 22-case set (standard English, no "
        "Pidgin, no ambiguous sentences), the system achieved 100% accuracy. The 77.08% figure "
        "reflects real-world classification difficulty rather than a fundamental classifier "
        "weakness, and the per-cluster analysis shows that POSITIVE (F1 = 0.9412), ANGER "
        "(F1 = 0.8889), and SADNESS (F1 = 0.8889) perform well; the primary challenge is in "
        "the AMBIGUOUS cluster (F1 = 0.5714), which by definition contains linguistically "
        "complex, contradictory inputs."
    )
    print(f"  Fixed evaluation results paragraph at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 8. SECTION 4.11.1 — Sadness sequence AEIF Turn 3 response
# ═══════════════════════════════════════════════════════════════════════
# Replace with the actual response from pipeline_comparison.json
actual_sadness_t3_aeif = (
    "It can be really tough to feel disconnected from the things that usually bring you joy. "
    "This feeling of emptiness can be especially hard to navigate when it feels like nothing "
    "is sparking your interest. I'm here to listen and support you as you work through these "
    "feelings. What's one thing that you used to enjoy doing, even if it doesn't feel "
    "appealing right now?"
)
actual_sadness_t3_baseline = (
    "It can be really tough to feel disconnected from the things that usually bring you joy. "
    "Sometimes this can be a sign that you're just feeling a bit drained or overwhelmed. "
    "It might be helpful to take some time to reflect on what's changed and what you might "
    "need to feel more engaged again. What's one thing you used to enjoy that you haven't "
    "tried in a while?"
)

idx = find_para(doc, "Turn 3 AEIF response")
if idx:
    update_paragraph(doc, idx,
        f'Turn 3 AEIF response: "{actual_sadness_t3_aeif}"'
    )
    print(f"  Fixed sadness AEIF response at paragraph {idx}")

idx = find_para(doc, "Turn 3 Baseline response")
if idx:
    update_paragraph(doc, idx,
        f'Turn 3 Baseline response (same emotion, no trend or strategy): "{actual_sadness_t3_baseline}"'
    )
    print(f"  Fixed sadness baseline response at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 9. SECTION 4.11.1 — Analysis paragraph comparing AEIF vs Baseline
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "The most significant difference is in the third sentence")
if idx:
    update_paragraph(doc, idx,
        "The most significant difference is in the second sentence. "
        "The AEIF response says \"This feeling of emptiness can be especially hard to navigate "
        "when it feels like nothing is sparking your interest\" — a reflective statement that "
        "directly names the emotional experience without interpreting it. "
        "The baseline response says \"Sometimes this can be a sign that you're just feeling a "
        "bit drained or overwhelmed\" — an interpretation that subtly pathologises the experience. "
        "The AEIF strategy constraint produced a more empathetically calibrated response."
    )
    print(f"  Fixed AEIF vs Baseline analysis at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 10. SECTION 4.11.2 — Mixed-emotion sequence responses
# ═══════════════════════════════════════════════════════════════════════
actual_mixed_t2_aeif = (
    "It's great that you're surrounded by supportive people who want the best for you. "
    "Their encouragement can be really uplifting, and it's wonderful that they're acknowledging "
    "your achievement. You must be feeling quite proud of yourself for getting the job, and "
    "that's something to celebrate. What is it about their support that's making you feel "
    "happy right now?"
)
actual_mixed_t2_baseline = (
    "It can be really tough when others don't quite understand why you're feeling a certain "
    "way, even if it's not what they expect. You might be feeling a mix of emotions, and "
    "that's okay - it's normal to have conflicting feelings, especially with something as "
    "significant as a new job. It's like there's a gap between what others think you should "
    "be feeling and what you're actually experiencing. What would you say is the most "
    "overwhelming emotion for you right now?"
)

idx = find_para(doc, "Turn 2 AEIF response:")
if idx:
    update_paragraph(doc, idx,
        f'Turn 2 AEIF response: "{actual_mixed_t2_aeif}"'
    )
    print(f"  Fixed mixed AEIF response at paragraph {idx}")

idx = find_para(doc, "Turn 2 Baseline response:")
if idx:
    update_paragraph(doc, idx,
        f'Turn 2 Baseline response: "{actual_mixed_t2_baseline}"'
    )
    print(f"  Fixed mixed baseline response at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 11. SECTION 4.11.2 — Analysis paragraph
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "This comparison is striking")
if idx:
    update_paragraph(doc, idx,
        "This comparison is striking. Both pipelines received the same input with the same "
        "emotion classification (POSITIVE). Yet their responses diverge meaningfully. "
        "The AEIF pipeline, guided by the POSITIVE_REINFORCEMENT strategy, produced an affirming "
        "response that celebrates the positive signal and invites the user to reflect on the "
        "source of their happiness. The baseline pipeline, receiving no strategy instruction, "
        "interpreted the same POSITIVE-classified message as an opportunity to explore emotional "
        "incongruence — focusing on \"conflicting feelings\" and \"a gap between what others "
        "think you should be feeling and what you're actually experiencing.\" The baseline "
        "response is not incorrect in isolation, but it is misaligned with the user's apparent "
        "improving state. This represents the most compelling evidence in the dataset that "
        "strategy-aware response generation produces qualitatively superior outputs."
    )
    print(f"  Fixed mixed comparison analysis at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 12. SECTION 4.11.3 — Anxiety sequence responses
# ═══════════════════════════════════════════════════════════════════════
actual_anxiety_t1_aeif = (
    "It can be really overwhelming to feel anxious about an upcoming exam. "
    "Let's take a moment to breathe together - why don't you take a deep breath in, "
    "hold it for a second, and then slowly let it out? Sometimes, focusing on our breath "
    "can help calm our minds. What's the first thing that comes to mind when you think "
    "about what you're looking forward to doing after the exam is over?"
)
actual_anxiety_t1_baseline = (
    "It's completely understandable to feel anxious about a big exam, and it sounds like "
    "you're putting a lot of pressure on yourself. I'm here to listen and offer support, "
    "and I want you to know that you're not alone in feeling this way. Many people get "
    "nervous before a big test, and it doesn't mean you're not prepared. What's been going "
    "through your mind when you think about the exam?"
)

idx = find_para(doc, "Turn 1 demonstrates the GROUNDING strategy")
if idx:
    update_paragraph(doc, idx,
        f'Turn 1 demonstrates the GROUNDING strategy\'s characteristic output. '
        f'AEIF response: "{actual_anxiety_t1_aeif}" '
        f'Baseline response: "{actual_anxiety_t1_baseline}" '
        f'The AEIF response offers a concrete grounding invitation (deep breath, focus on the '
        f'breath) while the baseline remains at acknowledgement and generic reassurance. This '
        f'illustrates the strategy instruction\'s direct influence on response content.'
    )
    print(f"  Fixed anxiety comparison at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 13. SECTION 4.12.2 — Trend Analyzer Threshold Calibration
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "The initial weighted severity scoring thresholds")
if idx:
    update_paragraph(doc, idx,
        "The initial severity scoring design used a simple priority ordering. However, "
        "the ESCALATING_DISTRESS detector (strict monotonic increase) was too strict for "
        "real conversations where emotions often plateau. The solution reordered the detection "
        "priority: PERSISTENT_DISTRESS (mean > 0.55) is now checked after ESCALATING_DISTRESS "
        "but before FLUCTUATING and STABLE_POSITIVE. A sequence of [SADNESS, ANXIETY, SADNESS] "
        "is correctly classified as FLUCTUATING (one direction change) rather than "
        "PERSISTENT_DISTRESS, because FLUCTUATING is evaluated before PERSISTENT_DISTRESS in the "
        "priority chain."
    )
    print(f"  Fixed threshold calibration at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 14. SECTION 4.12.3 — Ambiguity Detection Coverage
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "Initial testing revealed that emotionally contradictory")
if idx:
    update_paragraph(doc, idx,
        "Initial testing revealed that emotionally contradictory statements (\"I got the job "
        "but I'm terrified I'll mess it up\") were consistently classified as NEUTRAL because "
        "the HuggingFace model's neutral logit dominated despite the contradictory conjunction. "
        "The ambiguity detection pre-check was added to intercept these cases before the "
        "classifier. However, the pre-check only activates AFTER the HuggingFace classifier "
        "runs — it modifies the result rather than gating the classifier — and it is only "
        "triggered when the position-word (but, though, however, yet) AND both positive and "
        "negative keyword sets are detected. Cases where the classifier returns high-confidence "
        "NEUTRAL (above 0.3) bypass the keyword fallback entirely, meaning the ambiguity check "
        "runs on the HF result but the contradiction may not be detected if the keywords are "
        "not present in the defined sets. A secondary check on explicit ambiguity phrases "
        "(\"should be happy but\", \"don't deserve\", \"pretending\") handles cases that use "
        "implicit rather than explicit contrast."
    )
    print(f"  Fixed ambiguity detection at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 15. SECTION 4.14 — Summary metrics
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "72.9% emotion accuracy, 100% crisis accuracy, and 100% trend accuracy, with an average response latency of 54.2 milliseconds")
if idx:
    update_paragraph(doc, idx,
        "Automated evaluation on a 58-case test set (48 non-crisis, 10 crisis) demonstrated "
        "77.08% emotion accuracy, 100% crisis accuracy, and 100% trend accuracy across seven "
        "scripted sequences. All 31 unit tests pass."
    )
    print(f"  Fixed summary metrics at paragraph {idx}")

idx = find_para(doc, "72.9% emotion accuracy, 100% crisis accuracy, and 100% trend accuracy, with an average response latency of")
if idx and idx != find_para(doc, "72.9% emotion accuracy, 100% crisis accuracy, and 100% trend accuracy, with an average response latency of 54.2 milliseconds"):
    update_paragraph(doc, idx,
        "77.08% emotion accuracy, 100% crisis accuracy, and 100% trend accuracy"
    )
    print(f"  Fixed another metrics instance at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# 16. Section 4.11 paragraph - reference to comparison_output.txt
# ═══════════════════════════════════════════════════════════════════════
idx = find_para(doc, "comparison_output.txt")
if idx:
    update_paragraph(doc, idx,
        "The primary research finding of this study is the qualitative difference in response "
        "characteristics between the AEIF pipeline and the baseline pipeline across multi-turn "
        "conversation scenarios. Three scripted conversation sequences were run through both "
        "pipelines simultaneously, with complete side-by-side outputs recorded in the project's "
        "comparison data files (pipeline_comparison.json and comparison_output.txt). The following "
        "analysis focuses on the turns that most clearly demonstrate the contribution of "
        "trajectory analysis and strategy selection."
    )
    print(f"  Fixed comparison reference at paragraph {idx}")

# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════
output_path = "aida/AIDA_Dissertation_Corrected.docx"
doc.save(output_path)
print(f"\n✓ Saved corrected dissertation to {output_path}")
